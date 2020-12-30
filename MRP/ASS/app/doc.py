# encording:utf-8
import pythoncom

import pymssql
import os
import pandas as pd
from datetime import datetime, timedelta
from openpyxl import Workbook
from openpyxl import load_workbook
from openpyxl.styles import Border, Side, Alignment
from openpyxl.styles import Font
from os import listdir  # 设置虚拟路径,以供将py文件所在路径的类型文件全部导入
import win32com
from docx import Document  # docx用来操作word.docx,document用来新建空白文档
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL  # 单元格对齐方式位置模块
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 段落,表格对齐方式位置模块
from docx.enum.text import WD_LINE_SPACING  # 设置段落的行间距
from docx.enum.section import WD_ORIENT  # 章节的方向 PORTRAIT 纵向　LANDSCAPE 横向
from docx.enum.section import WD_SECTION  # 增加新章节 NEW_PAGE 下一页
from docx.oxml.ns import qn  # 设置中文字体类型
from docx.shared import Cm, Pt  # 设置字体大小Pt磅值,设置行高Cm厘米
from win32com.client import Dispatch, constants
import matplotlib.pyplot as plt
import matplotlib
import time
import re
import zipfile

db_port = '1433'
db_host = '10.10.91.180'
db_user = 'sa'
db_pwd = 'XMG3-Rel.1'
db_name = 'JCIHistorianDB'
Room_pid = [('1132', 23, 15),
            ('1133', 24, 16),
            ('1134', 25, 17),
            ('1135', 26, 18),
            ('1142B', 28, 20),
            ('1142A', 29, 21),
            ('1141A', 30, 22),
            ('1140', 39, 31),
            ('1159', 40, 32),
            ('1158', 41, 33),
            ('1157', 42, 34),
            ('1156', 14, 9),
            ('1153', 43, 35),
            ('1152', 44, 36),
            ('1151', 45, 37),
            ('1248A', 123, 111),
            ('1249A', 124, 112),
            ('1250A', 125, 113),
            ('1251A', 126, 114),
            ('1252A', 162, 141),
            ('1253A', 163, 142),
            ('1254A', 164, 143),
            ('1255A', 165, 144),
            ('1241A', 47, 145),
            ('1242A', 166, 146),
            ('1243A', 167, 147),
            ('1244A', 97, 86),
            ('1245A', 98, 87),
            ('1246A', 109, 100),
            ('1247A', 110, 101),
            ('1262A', 48, 127),
            ('1263A', 138, 128),
            ('1264A', 139, 129),
            ('1265A', 140, 130),
            ('1266A', 192, 168),
            ('1267A', 193, 169),
            ('1268A', 194, 170),
            ('1256A', 195, 171),
            ('1257A', 196, 172),
            ('1258A', 197, 173),
            ('1259A', 198, 174),
            ('1260A', 199, 175),
            ('1261A', 82, 75),
            ('1332', 247, 231),
            ('1333', 248, 232),
            ('1336', 249, 233),
            ('1342', 250, 234),
            ('1341', 251, 235),
            ('1340', 252, 236),
            ('1339', 253, 237),
            ('1338', 254, 238),
            ('1359', 268, 255),
            ('1358', 269, 256),
            ('1357', 270, 257),
            ('1356', 271, 258),
            ('1355', 272, 259),
            ('1352A', 273, 260),
            ('1351A', 274, 261),
            ('1432', 312, 300),
            ('1433', 313, 301),
            ('1434', 314, 302),
            ('1444', 316, 304),
            ('1443', 354, 335),
            ('1442A', 355, 336),
            ('1441', 356, 337),
            ('1440', 357, 338),
            ('1465', 329, 318),
            ('1464', 330, 319),
            ('1463', 331, 320),
            ('1462', 332, 321),
            ('1461', 333, 322),
            ('1460', 334, 323),
            ('1459', 358, 339),
            ('1456A', 359, 340),
            ('1455A', 360, 341),
            ('1454A', 361, 342),
            ('1453', 362, 343)]
roomid = [{'id':i, 'text': Room_pid[i][0] } for i in range(len(Room_pid))]
Alig_center = Alignment(horizontal='center', vertical='center',
                        wrapText=True)  # 双剧中
'''单元格线框颜色设定'''
border_All = Border(left=Side(border_style='thin', color='FF000000'),
                    right=Side(border_style='thin', color='FF000000'),
                    top=Side(border_style='thin', color='FF000000'),
                    bottom=Side(border_style='thin', color='FF000000'),
                    )


# 修改页眉
def docx_replace(old_file,new_file,rep):
    zin = zipfile.ZipFile(old_file, 'r')
    zout = zipfile.ZipFile(new_file, 'w')
    for item in zin.infolist():
        #print(item)
        b = zin.read(item.filename)
        if (item.filename ==  'word/header1.xml' in item.filename):
            #print('yes')
            res = b.decode("utf-8")
            for r in rep:
                res = res.replace(r,rep[r])
            b = res.encode("utf-8")
        else:
            #print('no')
            pass
        zout.writestr(item, b)
    zout.close()
    zin.close()


class RemoteWord:
    def __init__(self, filename=None):
        self.xlApp = win32com.client.DispatchEx('Word.Application')
        self.xlApp.Visible = 0
        self.xlApp.DisplayAlerts = 0  # 后台运行，不显示，不警告
        if filename:
            self.filename = filename
            if os.path.exists(self.filename):
                self.doc = self.xlApp.Documents.Open(filename)
            else:
                self.doc = self.xlApp.Documents.Add()  # 创建新的文档
                self.doc.SaveAs(filename)
        else:
            self.doc = self.xlApp.Documents.Add()
            self.filename = ''

    def add_doc_end(self, string):
        '''在文档末尾添加内容'''
        rangee = self.doc.Range()
        rangee.InsertAfter('\n' + string)

    def add_doc_start(self, string):
        '''在文档开头添加内容'''
        rangee = self.doc.Range(0, 0)
        rangee.InsertBefore(string + '\n')

    def insert_doc(self, insertPos, string):
        '''在文档insertPos位置添加内容'''
        rangee = self.doc.Range(0, insertPos)
        if (insertPos == 0):
            rangee.InsertAfter(string)
        else:
            rangee.InsertAfter('\n' + string)

    def replace_doc(self, string, new_string):
        '''替换文字'''
        self.xlApp.Selection.Find.ClearFormatting()
        self.xlApp.Selection.Find.Replacement.ClearFormatting()
        self.xlApp.Selection.Find.Execute(string, False, False, False, False, False, True, 1, True, new_string, 2)

    def save(self):
        '''保存文档'''
        self.doc.Save()

    def save_as(self, filename):
        '''文档另存为'''
        self.doc.SaveAs(filename)

    def close(self):
        '''保存文件、关闭文件'''
        self.save()
        self.xlApp.Documents.Close()
        self.xlApp.Quit()


class SqlServerOperate(object):

    def __init__(self, server, port, user, password, db_name, as_dict=False):
        self.server = server
        self.port = port
        self.user = user
        self.password = password
        self.db_name = db_name
        self.conn = self.get_connect(as_dict=as_dict)
        pass

    def __del__(self):
        self.conn.close()

    def get_connect(self, as_dict=False):
        conn = pymssql.connect(
            server=self.server,
            port=self.port,
            user=self.user,
            password=self.password,
            database=self.db_name,
            as_dict=as_dict,
            charset="utf8"
        )
        return conn

    def exec_query(self, sql):
        cur = self.conn.cursor()
        cur.execute(sql)
        result_list = list(cur.fetchall())
        cur.close()

        # 使用with语句（上下文管理器）来省去显式的调用close方法关闭连接和游标
        # print('****************使用 with 语句******************')
        # with self.get_connect() as cur:
        #     cur.execute(sql)
        #     result_list = list(cur.fetchall())   # 把游标执行后的结果转换成 list
        #     # print(result_list)

        return result_list


def sql_query(wpsid, data0, data1):
    '''
    根据pid 开始结束时间返回一个元组列表
    '''
    ms = SqlServerOperate(db_host, db_port, db_user, db_pwd, db_name)
    # sql_string = "SELECT PointName, PointID FROM dbo.RawDigital WHERE (PointName LIKE '%%')"
    sql_string = '''SELECT PointSliceID, UTCDateTime, ActualValue 
                    FROM FloatSampleView
                    WHERE (UTCDateTime >= '{}')
                    AND (UTCDateTime < '{}')
                    AND (PointSliceID = '{}')
                    '''.format(data0, data1, wpsid)
    # print(sql_string)
    sql_string2 = '''SELECT PointSliceID, UTCDateTime, OtherValue 
                    FROM tblOtherValueFloat
                    WHERE (UTCDateTime >= '{}')
                    AND (UTCDateTime < '{}')
                    AND (PointSliceID = '{}')
                    '''.format(data0, data1, wpsid) # 20191101
    temp_result_list = ms.exec_query(sql_string) + ms.exec_query(sql_string2) # 20191101
    temp_result_list = sorted(temp_result_list, key=lambda x: x[1]) # 20191101
    return temp_result_list


def room_to_id(room_id):
    '''
    房间号匹配温度湿度pid无法匹配返回False
    '''
    for x, y, z in Room_pid:
        if room_id == x:
            wd_pid = y
            sd_pid = z
            if wd_pid or sd_pid:
                return wd_pid, sd_pid


def wsd_query(wd_pid, sd_pid, data0, data1):
    '''
    输入:开始结束时间及房间号码
    返回:修正格式后温度湿度列表
    '''
    wd_list = sql_query(wpsid=wd_pid, data0=data0, data1=data1)
    sd_list = sql_query(wpsid=sd_pid, data0=data0, data1=data1)
    # 时间加8小时系统默认记录0时区时间.UTC +08:00

    wd_list = [(x, '{} CST'.format((y + timedelta(hours=8)).strftime("%m/%d/%y %I:%M %p")), '{:.1f}'.format(float(z))) for x, y, z in wd_list if z != None] # 20191101
    sd_list = [(x, '{} CST'.format((y + timedelta(hours=8)).strftime("%m/%d/%y %I:%M %p")), '{:.1f}'.format(float(z))) for x, y, z in sd_list if z != None] # 20191101
    if len(wd_list) % 96 == 0:
        return wd_list, sd_list
    else:
        print('数据异常不等于96倍数请检查')
        return wd_list, sd_list


def datetformat(dateinput):
    '''
    输入:时间字符串
    输出:格式化后的时间格式转换为可以调用标准字符-8时区
    '''
    date_out = datetime.strptime(dateinput, '%Y%m%d')
    date_out = date_out - timedelta(hours=8)
    return date_out


def exceed_out_list_w(room_id, rew_list):
    '''输入房间编号温度列表输出异常数据列表'''
    '''根据ID返回范围'''
    room_id_2 = int(room_id[1:2])  # 根据房间编号判断是第几层
    if room_id in ['1151', '1152', '1153']:
        er_l = [(x, y, z) for x, y, z in rew_list if float(z) > float(29) or float(z) < float(18)]
        return er_l
    elif room_id_2 == 2:
        er_l = [(x, y, z) for x, y, z in rew_list if float(z) > float(26) or float(z) < float(20)]
        return er_l
    elif room_id_2 in [1, 3, 4]:
        er_l = [(x, y, z) for x, y, z in rew_list if float(z) > float(26) or float(z) < float(16)]
        return er_l


def exceed_out_list_s(res_list):
    '''输入湿度列表导出异常数据列表'''
    er_l = [(x, y, z) for x, y, z in res_list if float(z) > float(70) or float(z) < float(40)]
    return er_l


def exceed_out_list_w_day(rew_list):
    '''输入温度列表返回日温差最大最小,以及异常数据列表.并返回一个超标天数天数'''
    num96 = len(rew_list) / 96
    xx = [0] + [a + 1 for a in [i * 96 for i in range(1, int(num96))]]  # 生成一个从0 开始到96 倍数用于切片的列表
    range_s_t = [(xx[i - 1], xx[i]) for i in range(1, len(xx))]  # 组装切片列表[(0, 97)(97, ...)]
    wen_cha_day = []  # 每天的最大值减去最小值的值
    exceed_day_value = []  # 超标异常数据列表
    exceed_day_count = 0  # 超标天数
    for k in range_s_t:
        rew_day = rew_list[k[0]:k[1]]
        f_temp = float(max(rew_day, key=lambda n: n[2])[2]) - float(min(rew_day, key=lambda n: n[2])[2])
        wen_cha_day.append(f_temp)
        if f_temp > 4:
            exceed_day_value.append(rew_day)
            exceed_day_count += 1
    wenc_minmax = '{:.1f}-{:.1f}'.format(min(wen_cha_day), max(wen_cha_day))  # 计算一段时间内的最大最小值罗列
    return wenc_minmax, exceed_day_value, exceed_day_count


def creat_image(room_id, rew_list, res_list):
    '''输入房间编号温湿度列表输出图片'''
    room_title_wd = '{} 温度趋势'.format(room_id)
    room_title_sd = '{} 湿度趋势'.format(room_id)
    pngname = ['rpngwd.png', 'rpngsd.png']
    wy = [float(i[2]) for i in rew_list]
    wx = [i for i in range(1, len(rew_list)+1)]
    sy = [float(i[2]) for i in res_list]
    sx = [i for i in range(1, len(res_list)+1)]
    font = {'family': 'MicroSoft Yahei',
            'weight': 'bold',
            'size': 12}
    matplotlib.rc("font", **font)
    '''温度出图'''
    plt.figure(figsize=(9, 6))
    plt.grid(linestyle='-.', axis='y')
    plt.plot(wy, label=room_title_wd, linewidth=1, color='blue', marker='o', markerfacecolor='red', markersize=1)
    plt.title(room_title_wd, fontsize=24)
    plt.xlabel('点位', fontsize=12)
    plt.ylabel('温度°C', fontsize=12)
    plt.legend(loc='upper left')  # 设置数字标签
    # 调整x轴刻度
    _x = list(wx)[::len(wx)//20]
    _y = [float(i) for i in range(15, 31)]
    # 取步长，数字和字符串一一对应，数据长度一样
    plt.xticks(_x, rotation=45)  # rotation 旋转角度
    plt.yticks(_y)
    plt.savefig(pngname[0], dpi=96, bbox_inches='tight')
    '''湿度出图'''
    plt.figure(figsize=(9, 6))
    plt.grid(linestyle='-.', axis='y')
    plt.plot(sy, label=room_title_sd, linewidth=1, color='blue', marker='o', markerfacecolor='red', markersize=1)
    plt.title(room_title_sd, fontsize=24)
    plt.xlabel('点位', fontsize=12)
    plt.ylabel('湿度 RH%', fontsize=12)
    plt.legend(loc='upper left')  # 设置数字标签
    # 调整x轴刻度
    _x = list(sx)[::len(sx)//20]
    _y = [float(i) for i in range(0, 101, 5)]
    # 取步长，数字和字符串一一对应，数据长度一样
    plt.xticks(_x, rotation=45)  # rotation 旋转角度
    plt.yticks(_y)
    plt.savefig(pngname[1], dpi=96, bbox_inches='tight')
    return pngname


def room_to_wsdh(room_id):
    '''根据ID返回范围'''
    wsdh_1626 = '16 - 26 °C'  # 兔狗猴 1F 3F 4F
    wsdh_1829 = '18 - 29 °C'  # 豚鼠
    wsdh_2026 = '20 - 26 °C'  # SPF 2F
    room_id_2 = int(room_id[1:2])  # 根据房间编号判断是第几层
    if room_id in ['1151', '1152', '1153']:
        return wsdh_1829
    elif room_id_2 == 2:
        return wsdh_2026
    elif room_id_2 in [1, 3, 4]:
        return wsdh_1626


def Creat_docx(study_number, study_name, study_anmail, room_id, s_datetime, t_datetime, rep):
    pythoncom.CoInitialize()
    word_path_moban = r'K:\mashuaifei\偏执狂的Django\MRP\ASS\media\moban.docx'
    doc = Document(word_path_moban)
    s_wsdch = '{%study-wsdch%}'
    s_name = '{%study-name%}'
    s_num = '{%study-num%}'
    s_anmail = '{%study-anmail%}'
    s_room = '{%study-room%}'
    s_date_room = '{%study-date-room%}'

    wdfb = '试验期间该动物房温度分布情况见附图。'
    sdfb = '试验期间该动物房湿度度分布情况见附图。'
    wdfbday = '试验期间该动物房日温差异常数据见附件。'
    doc_t2 = doc.tables[2]  # 温度表
    doc_t3 = doc.tables[3]  # 日温差表
    doc_t4 = doc.tables[4]  # 湿度表
    s_t_date_room_t = []   # 日期列表
    exceed_sd_list_total = []  # 湿度总异常
    exceed_wd_list_total = []  # 温度总异常
    exceed_day_value_total = []  # 日温差总异常
    room_id = [roomid[int(i)].get('text') for i in room_id]
    print(room_id)
    for i in range(len(room_id)):
        print(room_id[i])
        wd_pid, sd_pid = room_to_id(room_id[i])
        print(s_datetime)
        data0, data1 = datetformat(s_datetime[i]), datetformat(t_datetime[i]) + timedelta(hours=24)  # 加24小时才能与时间对应包含输入的那一天.

        rew_list, res_list = wsd_query(wd_pid, sd_pid, data0, data1)
        pngname = creat_image(room_id[i], rew_list, res_list)  # 生成图片
        time_total = []
        exceed_sd_list_to = []  #每次循环异常列表
        exceed_wd_list_to = []
        exceed_day_value_to = []
        s_datetime_new = '{}-{}-{}'.format(s_datetime[i][:4], s_datetime[i][4:6], s_datetime[i][6:])
        t_datetime_new = '{}-{}-{}'.format(t_datetime[i][:4], t_datetime[i][4:6], t_datetime[i][6:])
        time_total.append(s_datetime_new)
        time_total.append(t_datetime_new)
        #print(time_total)  # ['2017-09-22','2017-10-11']
        s_t_date_room = '{} 至 {} ({})'.format(s_datetime_new, t_datetime_new, room_id[i])  # 隐患
        s_t_date_room_t.append(s_t_date_room)

        wd_min_max = '{}-{}'.format(min(rew_list, key=lambda n:n[2])[2], max(rew_list, key=lambda n:n[2])[2])
        sd_min_max = '{}-{}'.format(min(res_list, key=lambda n:n[2])[2], max(res_list, key=lambda n:n[2])[2])
        wd_count = '{}'.format(len(rew_list))
        sd_count = '{}'.format(len(res_list))

        '''温度超标列表及计数'''
        exceed_wd_list = exceed_out_list_w(room_id[i], rew_list)  # 异常列表，列表套元组[(dsadasd),(ddasdaseqw)]

        wd_exceed_count = '0'
        if exceed_wd_list:
            wd_exceed_count = '{}'.format(len(exceed_wd_list))
            exceed_wd_list_to.append(room_id[i])  # 先添加房间
            exceed_wd_list_to.append(time_total)  # 再添加时间
            exceed_wd_list_to.append(exceed_wd_list)  # 最后添加异常列表  # ['A1245',['time','time'],[(a),(b)]]
            exceed_wd_list_total.append(exceed_wd_list_to)  # 添加到总的里  #[ ['A1245',['time','time'],[(a),(b)]]， ['A1233',['time2','time4'],[(ds),(das)]] ]
        '''湿度超标列表及计数'''
        sd_exceed_count = '0'
        exceed_sd_list = exceed_out_list_s(res_list)
        if exceed_sd_list:
            #print(exceed_sd_list)
            exceed_sd_list_to.append(room_id[i])  # 先添加房间
            exceed_sd_list_to.append(time_total)  # 再添加时间
            exceed_sd_list_to.append(exceed_sd_list)  # 最后添加异常列表  # ['A1245',['time','time'],[(a),(b)]]
            exceed_sd_list_total.append(exceed_sd_list_to)  # 添加到总的里  #[ ['A1245',['time','time'],[(a),(b)]]， ['A1233',['time2','time4'],[(ds),(das)]] ]
            sd_exceed_count = '{}'.format(len(exceed_sd_list))

        wd_pass = '{}'.format(str((int(wd_count) - int(wd_exceed_count)) / int(wd_count) * 100)[:4])
        sd_pass = '{}'.format(str((int(sd_count) - int(sd_exceed_count)) / int(sd_count) * 100)[:4])
        wd_row = [room_id[i], wd_min_max, wd_count, wd_exceed_count, wd_pass]  # 温度数据表内容组装
        wenc_minmax, exceed_day_value, exceed_day_count = exceed_out_list_w_day(rew_list)
        if exceed_day_value:
            exceed_day_value_to.append(room_id[i])  # 先添加房间
            exceed_day_value_to.append(time_total)  # 再添加时间
            exceed_day_value_to.append(exceed_day_value)  # 最后添加异常列表  # ['A1245',['time','time'],[(a),(b)]]
            exceed_day_value_total.append(exceed_day_value_to)   #  添加到总的里  #[ ['A1245',['time','time'],[(a),(b)]]， ['A1233',['time2','time4'],[(ds),(das)]] ]
        wd_row_day = [room_id[i], wenc_minmax, exceed_day_count]  # 日温差数据表组装
        sd_row = [room_id[i], sd_min_max, sd_count, sd_exceed_count, sd_pass]  # 湿度数据表内容组装

        '''表2 温度记录部分 '''
        doc_t2.add_row()
        index_numt2 = len(doc_t2.rows) - 1
        for j in range(5):
            doc_t2.cell(index_numt2, j).text = wd_row[j]
        '''表3 温度记录部分 '''
        doc_t3.add_row()
        index_numt3 = len(doc_t3.rows) - 1
        for j in range(3):
            doc_t3.cell(index_numt3, j).text = str(wd_row_day[j])
        # 判断是否有异常数据，放在循环里似乎会有问题
        if int(wd_row_day[2]) != 0:
            doc_t3.add_row()
            rownum3 = len(doc_t3.rows) - 1
            doc_t3.cell(rownum3, 0).text = wdfbday
            for number in range(1, 3):  # 创建合并列表最后一行
                # 合并单元格 参数 table.cell(1,2).merge(table.cell(2,2))
                doc_t3.cell(rownum3, 0).merge(doc_t3.cell(rownum3, number))
        '''表4 温度记录部分 '''
        doc_t4.add_row()
        index_numt4 = len(doc_t4.rows) - 1
        for j in range(5):
            doc_t4.cell(index_numt4, j).text = sd_row[j]

        for k in pngname:  # 增加图片
            doc.add_picture(k, width=Cm(15.20))

    '''添加异常附表'''
    #print(exceed_day_value_total)
    #print(exceed_wd_list_total)
    #print(exceed_sd_list_total)


    if exceed_wd_list_total + exceed_sd_list_total + exceed_day_value_total:
        fp = doc.add_paragraph()
        fp.add_run('附表：').bold = True
    if exceed_wd_list_total:
        print('温度异常')
        for i in range(len(exceed_wd_list_total)):
            fpwd = doc.add_paragraph()
            fpwd.add_run('{}温度异常数据({}至{})'.format(exceed_wd_list_total[i][0], exceed_wd_list_total[i][1][0],
                                                  exceed_wd_list_total[i][1][1])).bold = True
            wdextable = doc.add_table(len(exceed_wd_list_total[i][2]) + 1, 2, style='cti-table')
            wdextable.cell(0, 0).text = 'Time'
            wdextable.cell(0, 1).text = '{}温度趋势'.format(room_id[i])
            for y, m in enumerate(range(1, len(exceed_wd_list_total[i][2]) + 1)):
                wdextable.cell(m, 0).text = exceed_wd_list_total[i][2][y][1]
                wdextable.cell(m, 1).text = exceed_wd_list_total[i][2][y][2]
    if exceed_sd_list_total:
        print('湿度异常')
        for i in range(len(exceed_sd_list_total)):
            fpsd = doc.add_paragraph()
            fpsd.add_run('{}湿度异常数据({}至{})'.format(exceed_sd_list_total[i][0], exceed_sd_list_total[i][1][0],
                                                  exceed_sd_list_total[i][1][1])).bold = True
            sdextable = doc.add_table(len(exceed_sd_list_total[i][2]) + 1, 2, style='cti-table')
            sdextable.cell(0, 0).text = 'Time'
            sdextable.cell(0, 1).text = '{}湿度趋势'.format(room_id[i])
            for y, m in enumerate(range(1, len(exceed_sd_list_total[i][2]) + 1)):
                sdextable.cell(m, 0).text = exceed_sd_list_total[i][2][y][1]
                sdextable.cell(m, 1).text = exceed_sd_list_total[i][2][y][2]

    if exceed_day_value_total:
        print('日温差异常')
        for i in range(len(exceed_day_value_total)):
            fpday = doc.add_paragraph()
            fpday.add_run('{}日温差异常数据({}至{})'.format(exceed_day_value_total[i][0], exceed_day_value_total[i][1][0],
                                                    exceed_day_value_total[i][1][1])).bold = True
            dayextable = doc.add_table(len(exceed_day_value_total[i][2]) + 1, 2, style='cti-table')
            dayextable.cell(0, 0).text = 'Time'
            dayextable.cell(0, 1).text = '{}温度趋势'.format(room_id[i])
            for y, m in enumerate(range(1, len(exceed_day_value_total[i][2]) + 1)):
                dayextable.cell(m, 0).text = exceed_day_value_total[i][2][y][1]
                dayextable.cell(m, 1).text = exceed_day_value_total[i][2][y][2]

    '''表2，4固定的注释'''
    doc_t2.add_row()
    rownum = len(doc_t2.rows) - 1
    doc_t2.cell(rownum, 0).text = wdfb
    for number in range(1, 5):  # 创建合并列表最后一行
        # 合并单元格 参数 table.cell(1,2).merge(table.cell(2,2))
        doc_t2.cell(rownum, 0).merge(doc_t2.cell(rownum, number))
    doc_t4.add_row()
    rownum4 = len(doc_t4.rows) - 1
    doc_t4.cell(rownum4, 0).text = sdfb
    for number in range(1, 5):  # 创建合并列表最后一行
        # 合并单元格 参数 table.cell(1,2).merge(table.cell(2,2))
        doc_t4.cell(rownum4, 0).merge(doc_t4.cell(rownum4, number))

    '''数据居中'''
    table = doc.tables
    for oTable in table:
        rows_num = len(oTable.rows)
        columns_num = len(oTable.columns)
        for i in range(columns_num):
            for j in range(rows_num):
                oTable.cell(j, i).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    '''替换列表'''
    room_id_total = '、'.join(room_id)
    s_t_date_room_total = '      '.join(s_t_date_room_t)
    replace_list = [(s_num, study_number),
                    (s_name, study_name),
                    (s_anmail, study_anmail),
                    (s_room, room_id_total),
                    (s_date_room, s_t_date_room_total),
                    (s_wsdch, room_to_wsdh(room_id[0])),
                    ]
    new_docx = study_number + '.docx'
    doc.save(new_docx)
    '''替换word中的名字部分'''
    namedocx = os.path.join(r'K:\mashuaifei\偏执狂的Django\MRP\ASS\\', new_docx)
    replace_docx = RemoteWord(namedocx)
    for i in replace_list:
        replace_docx.replace_doc(i[0], i[1])
    replace_docx.close()

    '''替换页眉'''
    old = namedocx
    new = r'K:\mashuaifei\偏执狂的Django\MRP\ASS\\'+ study_number + 'ok.docx'
    print(old)
    print(new)
    docx_replace(old, new, rep)




def word(a, b, c, d, e1, e2):
    context = {}
    context['roomid'] = roomid
    # study_number = input('请输入实验编号:')
    study_number = a
    rep = {'studynum': study_number}
    # study_name = input('请输入实验名称:')
    study_name = b
    study_anmail = c
    # room_id = input([房间ID:','房间ID:'])
    room_id = d
    # s_datetime = input('起始日期:举例(20190725)')
    s_datetime = e1
    # 2017 09 22至 2017 11 13
    # 2017 11 14至 2017 12 11
    # t_datetime = input('结束日期:举例(20190725)')
    t_datetime = e2
    Creat_docx(study_number, study_name, study_anmail, room_id, s_datetime, t_datetime, rep)
    #jangs_excel(study_number, room_id, s_datetime, t_datetime)
    return context
