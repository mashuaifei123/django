import pandas as pd
import docx
from docx import Document  # docx用来操作word.docx,document用来新建空白文档
import win32com.client as wc
import win32com
import re
import time
import os
import pythoncom
import zipfile
class RemoteWord:
    def __init__(self, filename=None):
        self.xlApp = win32com.client.DispatchEx('Word.Application')
        self.xlApp.Visible = 0
        self.xlApp.DisplayAlerts = 0  # 后台运行，不显示，不警告
        if filename:
            self.filename = filename
            if os.path.exists(self.filename):
                self.doc = self.xlApp.Documents.Open(filename)
            else:
                self.doc = self.xlApp.Documents.Add()  # 创建新的文档
                self.doc.SaveAs(filename)
        else:
            self.doc = self.xlApp.Documents.Add()
            self.filename = ''

    def add_doc_end(self, string):
        '''在文档末尾添加内容'''
        rangee = self.doc.Range()
        rangee.InsertAfter('\n' + string)

    def add_doc_start(self, string):
        '''在文档开头添加内容'''
        rangee = self.doc.Range(0, 0)
        rangee.InsertBefore(string + '\n')

    def insert_doc(self, insertPos, string):
        '''在文档insertPos位置添加内容'''
        rangee = self.doc.Range(0, insertPos)
        if (insertPos == 0):
            rangee.InsertAfter(string)
        else:
            rangee.InsertAfter('\n' + string)

    def replace_doc(self, string, new_string):
        '''替换文字'''
        self.xlApp.Selection.Find.ClearFormatting()
        self.xlApp.Selection.Find.Replacement.ClearFormatting()
        self.xlApp.Selection.Find.Execute(string, False, False, False, False, False, True, 1, True, new_string, 2)

    def save(self):
        '''保存文档'''
        self.doc.Save()

    def save_as(self, filename):
        '''文档另存为'''
        self.doc.SaveAs(filename)

    def close(self):
        '''保存文件、关闭文件'''
        self.save()
        self.xlApp.Documents.Close()
        self.xlApp.Quit()




# path="K:\\mashuaifei\\iacuc\\B2020040-G1014-01" # 文件夹绝对路径
# files=[]
# for file in os.listdir(path):
#     if file.endswith(".doc"):
#         files.append(path+file)
#
# # doc文件另存为docx
# print(files)
# word = wc.Dispatch("Word.Application")
# for file in files:
#     doc = word.Documents.Open(file)
#     doc.SaveAs("{}x".format(file), 12)
#     doc.Close()
# word.Quit()
# print("完成！")
# word = wc.Dispatch("Word.Application")
# doc = word.Documents.Open("K:\\mashuaifei\\iacuc\\B2020040-G1014-01\\B2020040-G1014-01 食蟹猴静脉注射给予8MW2311和NovoA001 重复给药毒性预试验方案-终.doc")
# # 上面的地方只能使用完整绝对地址，相对地址找不到文件，且，只能用“\\”，不能用“/”，哪怕加了 r 也不行，涉及到将反斜杠看成转义字符。
# doc.SaveAs(r"K:\\mashuaifei\\iacuc\\B2020040-G1014-01\\B2020040-G1014-01 食蟹猴静脉注射给予8MW2311和NovoA001 重复给药毒性预试验方案-终.docx", 12, False, "", True, "", False, False, False,
#            False)  # 转换后的文件,12代表转换后为docx文件
# # doc.SaveAs(r"F:\\***\\***\\appendDoc\\***.docx", 12)#或直接简写
# doc.Close
# word.Quit

path = r"K:\\mashuaifei\\iacuc\\B2020040-G1014-01\\B2020040-G1014-01 食蟹猴静脉注射给予8MW2311和NovoA001 重复给药毒性预试验方案-终.docx"
file = docx.Document(path)
#for p in file.paragraphs:
#    print(p.text)
data = {}
lines = [0 for i in range(len(file.paragraphs))]
tables = file.tables
k = 0

for paragraph in file.paragraphs:
    lines[k] = paragraph.text
    k = k+1

#print(lines)
def locaation(index,lines):
    locate = []
    for j in index:
        for i in range(len(lines)):
            lines[i] = lines[i].replace(' ', '')
            if lines[i] == j:
                locate.append(i)
    return locate

def choose(index,lines1):
    data = []
    for j in index:
        if isinstance(j, list):
            for i in range(len(lines1)):
                if lines1[i].startswith(j[0]) or lines1[i].startswith(j[1]) :
                    if '：' in lines1[i]:
                        data.append(lines1[i])
                    else:
                        data.append(lines1[i+1])
        else:
            for i in range(len(lines1)):
                if lines1[i].startswith(j):
                #if lines[i] == j:
                    if '：' in lines1[i]:
                        data.append(lines1[i])
                    else:
                        data.append(lines1[i+1])
    return data


def docx_replace(old_file,new_file,rep):
    zin = zipfile.ZipFile (old_file, 'r')
    zout = zipfile.ZipFile (new_file, 'w')
    for item in zin.infolist():
        buffer = zin.read(item.filename)
        if (item.filename == 'word/document.xml' or 'header' in item.filename):
            res = buffer.decode("utf-8")
            for r in rep:
                res = res.replace(r,rep[r])
            buffer = res.encode("utf-8")
        zout.writestr(item, buffer)
    zout.close()
    zin.close()

index_1 = ['专题名称及编号','试验材料']
index_2 = ['试验系统', '动物的饲养和管理']
index_3 = ['试验材料', '溶媒']
index_3_1 = ['溶媒', '试验系统']
index_4 = ['试验设计', '观察与检查']
index_5 = [ '观察与检查', '大体解剖、脏器称量和组织病理学检查']
#index_6 = [ '动物饲养']
index_6 = ['临床病理', '大体解剖、脏器称量和组织病理学检查']
index_7 = ['一般状态观察', '临床病理']
index_7_1 = ['临床病理','大体解剖、脏器称量和组织病理学检查']
index_8 = ['大体解剖、脏器称量和组织病理学检查','数据采集和分析']

index_1_locate = locaation(index_1, lines)
index_2_locate = locaation(index_2, lines)
index_3_locate = locaation(index_3, lines)
index_3_1_locate = locaation(index_3_1, lines)
index_4_locate = locaation(index_4, lines)
index_5_locate = locaation(index_5, lines)
if len(index_5_locate)==1:
    index_5 = ['观察与检查', '大体解剖观察和组织病理学检查']
    index_5_locate = locaation(index_5, lines)
index_6_locate = locaation(index_6, lines)
if len(index_6_locate)==1:
    index_6 = ['观察与检查', '大体解剖观察和组织病理学检查']
    index_6_locate = locaation(index_6, lines)
index_7_locate = locaation(index_7, lines)
index_7_1_locate = locaation(index_7_1, lines)
if len(index_7_1_locate)==1:
    index_7_1 = ['观察与检查', '大体解剖观察和组织病理学检查']
    index_7_1_locate = locaation(index_7_1, lines)
index_8_locate = locaation(index_8, lines)
if len(index_8_locate)==1:
    index_8 = index_8 = ['大体解剖观察和组织病理学检查','数据采集和分析']
    index_8_locate = locaation(index_8, lines)

print(index_1_locate)
print(index_2_locate)
print(index_3_locate)
print(index_4_locate)
print(index_5_locate)
print(index_6_locate)
print(index_7_locate)

lines1 = lines[index_1_locate[0]: index_1_locate[1]]

lines2 = lines[index_2_locate[0]: index_2_locate[1]]
lines3 = lines[index_3_locate[0]: index_3_locate[1]]
lines3_1 = lines[index_3_1_locate[0]: index_3_1_locate[1]]
lines4 = lines[index_4_locate[0]: index_4_locate[1]]
lines5 = lines[index_5_locate[0]: index_5_locate[1]]
lines6 = lines[index_6_locate[0]: index_6_locate[1]]
lines7 = lines[index_7_locate[0]: index_7_locate[1]]
lines7_1 = lines[index_7_1_locate[0]: index_7_1_locate[1]]
lines8 = lines[index_8_locate[0]: index_8_locate[1]]

index_basic = ['专题名称：', '专题编号：', '试验目的', '姓名：', '试验操作：']
index_basic1 = ['{%study-title%}', '{%study-num%}','{%study-purpose%}', '{%sd%}','{%sy%}','{%person%}','{%day1%}' ,'{%day2%}', '{%day3%}','{%day4%}']
index_basic_day4 = ['解剖时间']

index_animal = ['品系：', '等级：', '性别和数量','体重：', '年龄：','来源', '实验动物选择依据和数量','动物数量选择理由：']
index_animal1 = ['{%species%}', '{%level%}', '{%n1%}','{%n2%}', '{%age1%}','{%age2%}','{%weight1%}','{%weight2%}','{%sourse%}', '{%reason1%}','{%reason2%}','{%reason3%}','{%agency1%}','{%agency2%}']

index_medicine = ['名称/代号：', '名称：','代号：']
index_design = ['组别设计：', '动物数量：', '性别比例：', '给药剂量：', ['给药容积：', '给药体积：'], '给药途径：', '给药速率：', '给药频率及周期：']
index_procedure = ['{%substance1%}','{%substance2%}', '{%yang%}', '{%ying%}', '{%others%}','{%group1%}', '{%group2%}',  '{%group3%}',  '{%group4%}',
                  '{%group5%}', '{%animal_num1%}', '{%animal_num2%}', '{%animal_num3%}', '{%animal_num4%}', '{%animal_num5%}',
                    '{%route1%}', '{%route2%}',  '{%route3%}',  '{%route4%}',  '{%route5%}', '{%dosage1%}',  '{%dosage2%}',
                   '{%dosage3%}', '{%dosage4%}', '{%frequency%}','{%volume%}','{%rate%}'
                   ]


index_procedure1 = ['{%jiankang%}', '{%xueye%}', '{%niaoye%}', '{%mianyi%}','{%tk1%}','{%tk2%}','{%ada%}', '{%baoding%}', '{%method%}','{%siyang%}']

index_anlesi = ['{%anlesi%}','{%yanke%}','{%niaoye1%}']

index_specimens = ['{%catch_reason%}', '{catch_prostion%}', '{%catch_frequency%}', '{%catch_volume%}', '{%catch_maxvolume%}']

index_observe = ['一般状态观察','体重', '摄食量', '眼科检查', '免疫原性（ADA）', ' 毒代血样采集']
index_collect = ['一般状态观察','体重', '摄食量', '眼科检查', '免疫原性（ADA）', ' 毒代血样采集']



# 去除字符串的标点符号
punctuation = '!，;:?"\'；。（）'
def removePunctuation(text):
    text = re.sub(r'[{}]+'.format(punctuation), ' ', text)
    return text.strip()


def calculate(time1, time2, type='day'):
     day1 = time.strptime(str(time1), '%Y-%m-%d')
     day2 = time.strptime(str(time2), '%Y-%m-%d')
     if type == 'day':
      day_num = (int(time.mktime(day2)) - int(time.mktime(day1))) / (
       24 * 60 * 60)
     return abs(int(day_num))

# part1

data_basic = choose(index_basic, lines1)
data_basic1 = [removePunctuation(i.split('：')[1]) if '：' in i else i for i in data_basic]
data_basic1_1 = [removePunctuation(i.split('：')[0]) if '：' in i else i for i in data_basic]
xm=0
for i in data_basic1_1:
    if i.startswith('姓名'):
        xm+=1
if xm ==2:
    data_basic1.pop(4)
else:
    pass

data_basic11 = choose(index_basic_day4, lines8)
day4 = data_basic11[0]
if '，' in day4:
    day4 = day4.split('，')[0]
else:
    pass

try:
    index_day = ['预定试验开始日期：', '预定试验结束日期：', '预定首次给药日期：', '给药期结束解剖：']
    data_day = data_basic = choose(index_day, lines1)
    day_list = [removePunctuation(i.split('：')[1]) if '：' in i else i for i in data_basic]
    if len(day_list) ==4:
        day1 = calculate(day_list[0],day_list[1])
        day2 = calculate(day_list[0],day_list[2])
        day3 = calculate(day_list[2],day_list[3])
        day = [day1,day2,day3,day4]
    else:
        pass
except:
    index_basic_specail = [ '预定试验开始日期：', '预定试验结束日期：','预定雌性动物首次给药日期：', '预定雌性动物给药期结束解剖日期：']
    data_basic_special = choose(index_basic_specail, lines1)
    day_list = [removePunctuation(i.split('：')[1]) if '：' in i else i for i in data_basic_special]
    if len(day_list) ==4:
        day1 = calculate(day_list[0],day_list[1])
        day2 = calculate(day_list[0],day_list[2])
        day3 = calculate(day_list[2],day_list[3])
        day = [day1,day2,day3,day4]
    else:
        pass
    day1 = calculate(day_list[0],day_list[1])
    day2 = calculate(day_list[0],day_list[2])
    day3 = calculate(day_list[2],day_list[3])
    day = [day1,day2,day3,day4]

day = ['','','',day4]
part1 = data_basic1[:5] +['王海洋、徐堃、沈双琪、吴自强、付梦洁、吴若林、孟庆喜'] + day
print(part1)
print('part1 no problem')
#part1 finished

#part2
data_animal = choose(index_animal, lines2)
data_animal1 = [removePunctuation(i.split('：')[1]) if len(i) < 30  else i for i in data_animal ]
number =[int(int(re.findall("\d+",data_animal1[2] )[0])/2)]

if '约' in data_animal1[4]:
    age = [(data_animal1[4].split('分')[0] if '分' in data_animal1[4] else data_animal1[4]).split('约')[1]]
else:
    age = [data_animal1[4].split('时')[1]]

if  len(data_animal1[3]) < 10:
    weight = [data_animal1[3],data_animal1[3]]
elif '雌' in data_animal1[3]:
    # 体重：购入时雌性150 ~ 180 g、雄性160 ~ 190 g。
    weight =[(data_animal1[3].split('g')[0]).split('雌性')[1] + 'g', (data_animal1[3].split('g')[1]).split('雄性')[1] + 'g']
else:
    weight =[(data_animal1[3].split('g')[0]).split('时')[1] + 'g',(data_animal1[3].split('g')[0]).split('时')[1] + 'g']
sourse = [data_animal1[5]+'(或其他合格供应商)']
reason1 = [ data_animal1[6]]
reason2 = [data_animal1[7][:200]]
reason3 = [data_animal1[7][200:]]

agency_index = ['遵循的法规及技术指导原则', '质量保证']
index_agency_locate = locaation(agency_index, lines)

if len(index_agency_locate) == 1:
    agency_index1 = ['遵循的法规及技术指导原则', '试验时间安排']
    index_agency_locate = locaation(agency_index1, lines)
print(index_agency_locate)
lines_agency = lines[index_agency_locate[0]: index_agency_locate[1]]
agency_ = ['《药物重复', '《药物毒代']
agency = []
for j in agency_ :
    for i in lines_agency :
        if i.startswith(j):
            agency.append(i)
if len(agency) ==1:
    agency.append('')
part2 = data_animal1[:2] + number + number + age +age+ weight + sourse +reason1+reason2+reason3 +agency
print(part2)
print('part2 no problem')
 # part2 finished



#part3
data_medicine = choose(index_medicine, lines3)
data_medicine1 = [removePunctuation(i.split('：')[1]) if len(i)<40  else i for i in data_medicine ]
# 判断试验材料有哪些
'''
如果实验材料只有一个，必然是供试品；
如果实验材料有两个，且没有空白组，两个供试品：否则一个供试品，一个空白组；
如果三个，且有空白组，即两个供试品，一个空白组 如：A2020006-T014-01
新增： 空白组还可是叫辅料组
'''
test_materials = []
print(data_medicine1)

for i in range(len(lines3)):
    if lines3[i].startswith('名称/代号：') or lines3[i].startswith('名称：') or lines3[i].startswith('代号：'):
        test_materials.append(lines3[i-2])
print(test_materials)
if len(test_materials) == 1:
    substance1 = data_medicine1[0]
    substance2 = ''
elif len(test_materials) == 2 and '空白对照品' not in test_materials and '辅料对照品' not in test_materials:
    substance1 = data_medicine1[0]+'，'+data_medicine1[1]
    substance2 = data_medicine1[1]
elif len(test_materials) == 2 and '空白对照品' in test_materials or '辅料对照品' in test_materials:
    substance1 = data_medicine1[0]
    substance2 = data_medicine1[1]
else:
    substance1 = data_medicine1[0]+'，'+data_medicine1[1]
    substance2 = data_medicine1[2]


#判断溶媒，只有在没有空白组的情况下有
'''
溶媒的判断方法同试验材料
当溶媒只有一个，那必然唯一填入
当溶媒有两个，那必然会需要配制

'''
data_medicine_rongmei = choose(index_medicine, lines3_1)
data_medicine_rongmei1 = [removePunctuation(i.split('：')[1]) if len(i)<40  else i for i in data_medicine_rongmei ]

if '空白对照品' not in test_materials and '辅料对照品' not in test_materials:
    if len(data_medicine_rongmei) ==1:
        other = data_medicine_rongmei1[0]
    elif len(data_medicine_rongmei) ==2:
        for i in range(len(lines3_1)):
            if lines3_1[i].startswith('配制方法'):
                other = (lines3_1[i-1]).split('的')[0]
    data_medicine11 = [substance1, substance2, 'NA', 'NA', other]
else:
    other = substance2
    data_medicine11 = [substance1, substance2, 'NA', 'NA', other]
# if len(data_medicine1) > 2:
#     data_medicine11 = [data_medicine1[0]+'，'+data_medicine1[1], data_medicine1[2], 'NA, ', 'NA', data_medicine1[2]]
# elif len(data_medicine1) == 2:
#     data_medicine11 = [data_medicine1[0],data_medicine1[1], 'NA, ', 'NA', data_medicine1[1]]
# else :
#     data_medicine11 = [data_medicine1[0], data_medicine1[1],'NA, ', 'NA', 'NA']

# 组别的拆分
'''
组别的情况越来越多样了,新情况;  '组别设计：溶媒对照组、WXSH0376低剂量组、WXSH0376中剂量组和WXSH0376高剂量组；
主要需要区分“顿号”“逗号”“和”。
'''
data_design = choose(index_design, lines4)
print(data_design)
group1 = [[removePunctuation(data_design[0].split('：')[1])][0].split('、')[i] for i in range(len([removePunctuation(data_design[0].split('：')[1])][0].split('、')))]
for i in group1:
    if '和' in  i:
        group = [i.split('和')[0],i.split('和')[1]]
        group1.remove(i)
        if '中' in ''.join(group):
            group1 = group1 + group
        else:
            group1 = group + group1
    else:
        pass

# 动物数量
# print(data_design[1])
# print((re.findall(r'\d+', removePunctuation(data_design[1].split('：')[1]))))
if ' 'in data_design[1]:
    animal_num = str(int(int(re.findall(r'\d+',removePunctuation(data_design[1].split('：')[1]).split(' ')[-1])[0])/(len(group1)*2)))
else:
    animal_num = str(int(int(re.findall(r'\d+', removePunctuation(data_design[1].split('：')[1]))[1]) / (len(group1) * 2)))
animal_num = [animal_num + '             ' + animal_num] *len(group1)

# 给药途径的判断
'''
先给出大概的几种方式，找到给药途径的话语，判断这几种方式在不在其中。
# 部分试验方案会存在多种方式在其中，需要s数计算来判断
'''
route_choice = ['皮下', '静脉','灌胃']
s = 0
for i in route_choice:
    if i in data_design[5]:
        s+=1
if s>=2:
    route = ['']*len(group1)
elif '灌胃' in data_design[5]:
    route = ['灌胃' ] * len(group1)
else:
    route= [removePunctuation(data_design[5].split('：')[1])]*len(group1)


#通过表格找到给药剂量
'''
情况是 从表格中提取剂量，但是表格不固定为第几个
加一个判断，表格的第二列指标是否是"试验物质"
'''
for table in tables[6:7]:
    table_total = []
    for i, row in enumerate(table.rows[:]):   # 读每行
        row_content = []
        for cell in row.cells[:]:  # 读一行中的所有单元格
            c = cell.text
            row_content.append(c)
        table_total.append(row_content)
print(table_total)
if table_total[0][1] !='试验物质':
    for table in tables[5:6]:
        table_total = []
        for i, row in enumerate(table.rows[:]):   # 读每行
            row_content = []
            for cell in row.cells[:]:  # 读一行中的所有单元格
                c = cell.text
                row_content.append(c)
            table_total.append(row_content)

if '溶酶对照组'in group1 or '空白对照组' in group1 or '辅料对照组' in group1:
    table_total.pop(0)
    table_total.pop(0)
    table_total.pop(0)
else:
    table_total.pop(0)
    table_total.pop(0)
#print(table_total) #以列表形式导出每一行数据
dosage = [table_total[i][2] + 'mg/kg' for i in range(len(table_total))]

# 给药频率
'''
到底要不要写给药几周
'''
frequency = [(removePunctuation(data_design[-1].split('：')[1])).split('；')[0] ]

# 因为体积容积的区分

volume1 = [((removePunctuation(data_design[4].split('：')[1])).split(' ')[0]).split('为')[1] ]
if 'mL/kg' in volume1:
    volume = volume1
else :
    volume = [re.findall(r'\d+',removePunctuation(data_design[4].split('：')[1]))[0] +' mL/kg']
try:
    rate = [((removePunctuation(data_design[-2].split('：')[1]))).split('约')[1]]
except:
    rate = ['找不到给药速率NA']
if len(group1) < 5 :
    i = 5-len(group1)
    group1 = group1 + i*['']
    animal_num = animal_num + i*['']
    route = route + i*['']
    dosage = dosage + i*['']
else:
    group1 = group1
    animal_num = animal_num
    route = route
    dosage = dosage

part3 = data_medicine11 + group1 + animal_num +route + dosage + frequency +volume + rate
data_observe = choose(index_observe, lines5)
for table in tables[7:8]:
    for i, row in enumerate(table.rows[:]):   # 读每行
        row_content = []
        for cell in row.cells[:]:  # 读一行中的所有单元格
            c = cell.text
            row_content.append(c)
        #print (row_content) #以列表形式导出每一行数据
print(part3)
print('part3 no problem')


# part4
## 健康检查到临床病理
'''
首先是健康检查项目：先判断是否有没有健康检查项目 ，从zhibiao这个列表得知
                    如果有，即写成jiankang 
然后是临床检查项目： 判断血液，尿液是否在其中,从zhibiao1中得知

其他项目检查，同，判断ADA,TK在不在zhibiao1中
'''
sss = ['测定时间：', '观察时间：']
properties = []
zhibiao = []
for i in range(len(lines7)):
    if lines7[i].startswith('测定时间：') or lines7[i].startswith('观察时间：') or lines7[i].startswith('检测时间：') :
        property = lines7[i-1] + '：' + lines7[i]
        properties.append(property)
        zhibiao.append(lines7[i-1])

for i, j in enumerate(properties):
    if j.startswith('体重：'):
        properties[i] = ''.join(properties[i].split('（')[:-1])

# 体温，血压，心电，呼吸
if '体温' in ''.join(zhibiao):
    jiankang_1 = properties[3].replace('；','。')
    jiankang = ['体温、呼吸、心电、血压：'+(jiankang_1.split('时间：'))[1]]
else:
    jiankang = ['NA']

# 临床病理
properties_1 = []
zhibiao_1 = []
for i in range(len(lines7_1)):
    if lines7_1[i].startswith('测定时间：') or lines7_1[i].startswith('观察时间：')or lines7_1[i].startswith('检测时间：')  or lines7_1[i].startswith('检测的血样采集时间：')  or lines7_1[i].startswith('血样采集：') :
        property = lines7_1[i].split('：')[1][:-1]
        properties_1.append(property)
        zhibiao_1.append(lines7_1[i-1])
for i in range(len(lines7_1)):
    if lines7_1[i].startswith('采样时间：') or lines7_1[i].startswith('血样采集：'):
        if '剂量' in lines7_1[i]:
            if '采样时间：' in lines7_1[i]:
                property = lines7_1[i].split('采样时间：')[1] + lines7_1[i + 1]
            else:
                property = lines7_1[i].split('：')[1] + lines7_1[i + 1]
        else:
            property = lines7_1[i+1] + lines7_1[i + 2]
        properties_1.append(property)
        zhibiao_1.append(lines7_1[i-1])

print(zhibiao_1)
#尿液
if '尿液' in ''.join(zhibiao_1):
    niaoye1 = ['尿液：' + properties_1[1]]
    for i in range(len(lines7_1)):
        if '代谢笼' in lines7_1[i] or '底盘' in lines7_1[i]:
            niaoye_method1 = lines7_1[i]
    niaoye_method = niaoye_method1.split('：')[1]
    niaoye = [niaoye1[0]+ '（'+niaoye_method+'）。']
else:
    niaoye1 =['']
    niaoye = ['']
    niaoye_method = ''

# 血液
if '血液' in ''.join(zhibiao_1):
    xueye = ['血液学、凝血和血生化：' + properties_1[0]]
# 免疫
mianyi = ['免疫毒性指标：']

# 判断ADA,毒代是否在其中
if 'ADA' in ''.join(zhibiao_1):
    ada1 = properties_1[-2]
    ada = ['ADA：' + ada1 + '。']
else:
    ada1 = ''
    ada = ['']
if '毒代动力学' or '血样'in ''.join(zhibiao_1):
    tk_1 = properties_1[-1][0:200]
    tk_2 = properties_1[-1][200:]
else:
    tk_1 = ''
    tk_2 = ''
tk1 = ['TK：'+tk_1]
tk2 = [tk_2]

# 保定
index_bd =['动物的饲养和管理','试验设计']
index_bd_locate = locaation(index_bd, lines)
linesbd = lines[index_bd_locate[0]: index_bd_locate[1]]
print(linesbd)
for i in range(len(linesbd)):
    if linesbd[i].startswith('动物选择'):
        if '保定' in linesbd[i+2]:
            bd1 = linesbd[i+2]
        else:
            bd1 = linesbd[i+1]
if '不少于' in  bd1:
    j = re.sub('\D', '', bd1.split('不少于')[1])
else :
    j = 'NA'
if '猴' in data_animal1[0] and '保定' in bd1:
    i = '猴椅'
    baoding = '试验中涉及的保定方式有徒手保定和' + i + '保定，其中称重为徒手保定，保定时长约2 min；给药为'+i +'保定，保定时长为'+ j+'min。'
elif '鼠' in data_animal1[0] and '保定' in bd1 :
    i = '大鼠固定器'
    if j == 'NA' :
        baoding = '试验中涉及的保定方式有徒手保定和' + i + '保定，其中称重为徒手保定，保定时长约2 min；'
    else:
        baoding = '试验中涉及的保定方式有徒手保定和' + i + '保定，其中称重为徒手保定，保定时长约2 min；给药为'+i +'保定，保定时长为'+ j+'min。'

elif '犬' in data_animal1[0] and '保定' in bd1:
    i = '犬保定架'
    baoding = '试验中涉及的保定方式有徒手保定和' + i + '保定，其中称重为徒手保定，保定时长约2 min；给药为'+i +'保定，保定时长为'+ j+'min。'
else:
    baoding = '试验中涉及的保定方式有徒手保定保定，其中称重为徒手保定，保定时长约2 min。'

if '在' in bd1:
    baoding = [baoding + bd1.split('在')[1]]
else:
    baoding = [baoding]
for i in lines2:
    if i.startswith('剩余动物'):
        method = [i]
teminal = ['']

# 判断单笼合笼
for i in range(len(lines4)):
    if lines4[i].startswith('动物数量：'):
        shuliang =lines4[i]
    if lines4[i].startswith('性别比例：'):
        bili  = lines4[i]

if  int(re.sub('\D', '', shuliang.split('，')[0] )) >2:
    if '半' in bili:
        print('合笼')
        siyang = ['若动物出于兽医护理的需要或动物不相容的原因进行单笼饲养，将会作相应记录，并给予额外的环境丰富计划。']
else:
    print('单笼')
    siyang  =['']
part4 = jiankang + xueye + niaoye + mianyi + tk1 + tk2+ ada + baoding  + method +siyang

print(part4)
print('part4 no problem')
#part4 fininshed


#part other5
anlesi = []
for i in range(len(lines)):
    if lines[i] == '麻醉及安乐死方法':
        anlesi.append(lines[i+1])

if '眼科' in ''.join(zhibiao):
    c = '舒泰肌肉注射麻醉后进行检查，剂量参照机构SOP执行。'
else:
    c =''
anlesi.append(c)
anlesi.append(niaoye_method)



print(anlesi)
catch_huoti=[]
catch_huoti_index = []

# # 判断活体动物采集
# for i in range(len(lines7_1)):
#     if lines7_1[i].startswith('检测动物：') or lines7_1[i].startswith('采样动物：') or lines7_1[i].startswith('测定动物：') or lines7_1[i].startswith('检测的血样采集时间：') :
#         catch_huoti.append(lines7_1[i])
#         catch_huoti_index.append(i)  # 位置
# # 得到操作检测动物列表  ['检测动物：主试验各组待解剖大鼠；', '检测动物：主试验各组待解剖大鼠；', '检测动物：毒代卫星组存活动物；', '采样动物：所有计划测定的存活TK动物；']
# print(catch_huoti)
# print(catch_huoti_index)
#
# #判断,如果都有活体,执行if,否则sles
# if all("存活" in s for s in catch_huoti):
#     print('在临床就使用活体')
#     catch_reason = ['临床病理、TK、ADA血样']
#     for table in tables[7:8]:
#         row_content = []
#         for i, row in enumerate(table.rows[:]):  # 读每行
#             for cell in row.cells[:]:  # 读一行中的所有单元格
#                 c = cell.text
#                 #c = ''.join(filter(str.isdigit, c))
#                 c = ''.join(re.findall(r"\d+\.?\d*", c))
#                 row_content.append(c)
#     if row_content[0] =='':
#         for table in tables[8:9]:
#             row_content = []
#             for i, row in enumerate(table.rows[:]):  # 读每行
#                 for cell in row.cells[:]:  # 读一行中的所有单元格
#                     c = cell.text
#                     #c = ''.join(filter(str.isdigit, c))
#                     c = ''.join(re.findall(r"\d+\.?\d*", c))
#                     row_content.append(c)
#
#     row_content1 = [row_content[4] ,row_content[7] ,row_content[10]]
#     catch_volume = '临床病理样本约'+ str(sum([float(i) for i in row_content1])) + 'mL'
#     print(catch_volume)
#     catch_position = [
#         (([lines7_1[i] for i in range(len(lines7_1)) if lines7_1[i].startswith('采样方法：')][0]).split('，')[0]).split(
#             '：')[1]]
#     catch_volume_number = [float(''.join(re.findall(r"\d+\.?\d*", lines7_1[i]))) for i in range(len(lines7_1)) if
#                            lines7_1[i].startswith('采样方法：')]
#     catch_maxvolume_number = max(catch_volume_number)
#
#     catch_maxvolume = [catch_volume]
#     if isinstance(catch_maxvolume_number, float):
#         if len(catch_volume_number) == 2:
#             label = ['ADA', 'TK']
#             catch_maxvolume = [
#                 label[catch_volume_number.index(catch_maxvolume_number)] + '约' + str(catch_maxvolume_number) + 'mL']
#         else:
#             catch_volume_sum = ['找不到tk,ada其中一个的采血量']
#     else:
#         catch_volume_sum = ['找不到tk,ada的采血量']
# else:
#     # 由于临床没有活体 ,找到使用活体的指标
#     print('只有血样之后采用活体')
#     catch_reason = ['TK、ADA血样']
#     label = ['ADA','TK']
#     cunhuo_list = [catch_huoti_index[i]  for i in range(len(catch_huoti))  if '存活' in  catch_huoti[i] ]
#     catch_volume = [lines7_1[i] for i in cunhuo_list]
#     catch_position = [(([ lines7_1[i]  for i in range(len(lines7_1)) if lines7_1[i].startswith('采样方法：')][0]).split('，')[0] ).split('：')[1]]
#     catch_volume_number = [float(''.join(re.findall(r"\d+\.?\d*", lines7_1[i])))  for i in range(len(lines7_1)) if lines7_1[i].startswith('采样方法：') or  lines7_1[i].startswith('采样方法和采样量：') ]
#     catch_maxvolume_number = max(catch_volume_number)
#     catch_maxvolume =[label[catch_volume_number.index(catch_maxvolume_number)] + '约' + str(catch_maxvolume_number) +'mL']
#
#     catch_volume_sum_number = sum(catch_volume_number)
#     catch_volume_sum =['约' + str(catch_volume_sum_number) + 'mL' +'（' +catch_maxvolume[0] + '，'+ 'TK约' + str(catch_volume_number[1]) +'mL）']
# catch_frequency = ['24小时内不超过三次']
#
# print(catch_maxvolume)
# part_other = catch_reason +catch_position +catch_frequency +catch_maxvolume + catch_volume_sum
#finished

part_other = ['','','','','']
print(part1)
print(part2)
print(part3)
print(part4)
print(anlesi)
print(part_other )
# print(data_observe)





index = [0 for i in range(10)]
for j in index_1:
    for i in range(len(lines)):
        #if lines[i].startswith(j):
        if lines[i] == j:
            print(i)
            print(lines[i])

# pythoncom.CoInitialize()
# word_name = 'BTC-IAC-0011-3.0 IACUC-动物管理和使用审批表 3.0 2020-03-06生效'
# word_path_moban = r'K:\mashuaifei\iacuc\iacuc1\新建文件夹\moban.docx'
# doc = Document(word_path_moban)
# new_docx = part1[1] + word_name + '.docx'
# print(new_docx)
# doc.save(os.path.join(r'K:\mashuaifei\iacuc\iacuc1\新建文件夹\\', new_docx))
#
# replace_list = []
# for i,j in enumerate(index_basic1):
#     replace_list.append((j,part1[i]))
# for i,j in enumerate(index_animal1):
#     replace_list.append((j, part2[i]))
# for i,j in enumerate(index_procedure):
#     replace_list.append((j, part3[i]))
# for i,j in enumerate(index_procedure1):
#     replace_list.append((j, part4[i]))
# for i, j in enumerate(index_anlesi):
#     replace_list.append((j, anlesi[i]))
# # for i, j in enumerate(index_specimens ):
# #     replace_list.append((j, part_other[i]))
#
# print(replace_list)
#
# name_path = 'K:\mashuaifei\iacuc\iacuc1\新建文件夹' + "\\" + new_docx
# namedocx = os.path.abspath('K:\mashuaifei\iacuc\iacuc1\新建文件夹' + "\\" + new_docx)
# replace_docx = RemoteWord(namedocx)
# for i in replace_list:
#     if len(i)>200:
#         print(i)
# for i in replace_list:
#     try:
#         replace_docx.replace_doc(i[0], i[1])
#     except:
#         replace_docx.close()
#         print('its wrong')
# replace_docx.close()
# rep = {'{%study-num%}': data_basic1[1]}
# '''替换页眉'''
# new_docx1 = part1[1] + word_name + '1.docx'
# new = 'K:\mashuaifei\iacuc\iacuc1\新建文件夹' + "\\" + new_docx1
# docx_replace(name_path, new, rep)