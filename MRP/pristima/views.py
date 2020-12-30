from django.shortcuts import render
from django.shortcuts import HttpResponse
from django.shortcuts import HttpResponseRedirect
from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import get_object_or_404
from django.db.models import F, Q
from tools.PTSAPI import BW_AND_PA, get_BodyWeight, get_ProArticle, get_BW_Group, ProSchedule_Parameter  # 导入颜色板
from studys.models import studysinfo
import datetime
from django.contrib.auth.decorators import login_required  # 确认是否已经登入装装饰器
from django.contrib.auth.decorators import permission_required
import pandas as pd



def Dosage_BW(request):
    '''给药&体重联合查询 '''
    # 前段需要的专题列表
    studyno_list = studysinfo.objects.all().values("id","studyno")
    studyno_josn = [{'id':o['id'], 'text':o['studyno']} for o in studyno_list]

    PA_column = ['分组', '雌给药量', '雄给药量', '雌给药体积', '雄给药体积', '雌给药浓度','雄给药浓度','雌组体重', '雄组体重', '总组体重', '雌配制体积','雄配制体积','总配制体积']
    PA_columns = [{"title": x} for x in PA_column]  # 拼接为前端需要样式字典

    BW_column = ['动物编号', '体重', '单位', '称重日期']
    BW_columns = [{"title": x} for x in BW_column]  # 拼接为前端需要样式字典
    
    """ 合并判断查询部分"""
    q1 = Q()  # Q查询
    q1.connector = 'AND'              # 连接方式 默认AND 可以设定为OR
    con = Q()
    q2 = Q()
    q2.connector = 'OR' 

    if request.method == 'GET':
        Fstudyid = request.GET.get('q', None)  # 得到搜索关键词'studyid'
        
        if Fstudyid:  # 查询专题编号
            q1.children.append(('id', int(Fstudyid)))
            con.add(q1, 'AND')
            con.add(q2, 'AND')  # 合并查询
            adate = studysinfo.objects.filter(con)  # 拼接的Q查询返回Queryset order_by 排序

            """开始执行查询"""
            adate = adate.values('studyno')
            
            try:
                for i in adate:
                    studyno = i['studyno']
                GBW = get_BodyWeight(studyno) # 得到体重数据
                GPA = get_ProArticle(studyno) # 得到给药数据
                BW, BW_DF_G = get_BW_Group(GBW)  # 计算并整形体重数据
                try:
                    df_PA = BW_AND_PA(BW_DF_G, GPA)
                    MLSum = '{:.2f}'.format(df_PA['Summl/kg'].sum())
                    KGSum = '{:.2f}'.format(df_PA['BWSum'].sum())
                    df_PA.iloc[:,1:] = df_PA.iloc[:,1:].applymap(lambda x: '{:.2f}'.format(x)) # 格式化小数显示
                    df_PA = df_PA.values.tolist() # 给药清单
                except:
                    df_PA = []  # 失败给药清单为空
                df_BW = BW.values.tolist()  # 体重清单
                adate = {'studyno':studyno, 'MLSum':MLSum, 'KGSum':KGSum,}
            except:
                error = '查询失败: 请检查pristima系统中给药剂量设定.或最近一次体重未称重.如还未解决请联系管理员'
            try:
                context = {'studyno':studyno_josn, 'testdata':adate, 'PAtitle': PA_columns, 'PAd':df_PA, 'BWtitle':BW_columns, 'BWd':df_BW}  # 传入到render变量参数用于被脚本调用
            except:
                context = {'studyno':studyno_josn,'error':error}

            return render(request, 'Dosage_BW.html', context)

        else:
            context = {'studyno':studyno_josn}  # 传入到render变量参数用于被脚本调用
            return render(request, 'Dosage_BW.html', context)


def ScheduleParameter(request):
    '''专题计划及参数查询返回 临检申请单 '''
    # 前段需要的专题列表
    studyno_list = studysinfo.objects.all().values("id","studyno")
    studyno_josn = [{'id':o['id'], 'text':o['studyno']} for o in studyno_list]
    SP_column = ['日程日期','试验阶段','尿液or血样','仪器名称', '参数指标', '动物编号合集', '动物数量计数']
    SP_columns = [{"title": x} for x in SP_column]  # 拼接为前端需要样式字典

    """ 合并判断查询部分"""
    q1 = Q()  # Q查询
    q1.connector = 'AND'              # 连接方式 默认AND 可以设定为OR
    con = Q()
    q2 = Q()
    q2.connector = 'OR' 
    if request.method == 'GET':
        Fstudyid = request.GET.get('q', None)  # 得到搜索关键词'studyid' 
        if Fstudyid:  # 查询专题编号
            q1.children.append(('id', int(Fstudyid)))
            con.add(q1, 'AND')
            con.add(q2, 'AND')  # 合并查询
            adate = studysinfo.objects.filter(con)  # 拼接的Q查询返回Queryset order_by 排序
            """开始执行查询"""
            adate = adate.values('studyno')
            try:
                for i in adate:
                    studyno = i['studyno']
                df = ProSchedule_Parameter(studyno)
                df_SP = df.values.tolist() # 调用函数输出列表

                '''以下为生成word'''
                from tools.config import outputfilepath
                from docx import Document
                address_col_row = [(1, 3,'Hematology (ADVIA 2120i)'),(1, 4,'Coagulation (CA-7000)'),(1, 5,'Blood Biochemistry (cobas 6000)'),(1, 6,'Urine Analysis (AX-4030)'),(4,6, 'Urine Sediment'),(1, 8,'BD FACSCalibur')]  # 具体的word 位置
                study_numb_phase_word = [(7, 0,'phase'),(7, 1,'采集日期'),(4, 1,'动物数量'),(1, 11,'动物编号')]
                study_numb_phase_df = [(0, 1,'phase'),(0, 0,'采集日期'),(0, 6,'动物数量'),(0, 5,'动物编号')]
                for x, y in df.groupby(['Days','PhaseName']): # 日期阶段分组后写入word
                    doc = Document(r'tools\doctable\BTC-CLI-0001.docx') # 源文件 test.docx
                    doc.tables[0].cell(col_idx=1,row_idx=0).text = str(studyno) # 写入专题编号
                    for word_cell in range(len(study_numb_phase_word)):
                        doc.tables[0].cell(col_idx=study_numb_phase_word[word_cell][0],row_idx=study_numb_phase_word[word_cell][1]).text = str(y.iloc[study_numb_phase_df[word_cell][0],study_numb_phase_df[word_cell][1]]) # 写入阶段日期及动物编号
                    z = y.set_index('Panel') # 设定列为索引用于标签loc选择
                    for adcr in address_col_row:  # 尝试填充仪器设备指标错误就跳过
                        try:
                            doc.tables[0].cell(col_idx=adcr[0],row_idx=adcr[1]).text = str(z.loc[adcr[2],'AbbSum'])    # 第0 个表格按照位置替换原始框内数据
                        except:
                            pass
                    doc.save(outputfilepath + '\{}_{}_{}.docx'.format(studyno,x[0],x[1]))  # 保存
                adate = {'studyno':studyno, 'outputfilepath':outputfilepath}
            except:
                error = '查询失败: 请检查pristima系统试验计划设定是否具备各种信息如还未解决请联系管理员'
            try:
                context = {'studyno':studyno_josn, 'testdata':adate, 'SPtitle': SP_columns, 'SPd':df_SP}  # 传入到render变量参数用于被脚本调用
            except:
                context = {'studyno':studyno_josn,'error':error}

            return render(request, 'ScheduleParameter.html', context)

        else:
            context = {'studyno':studyno_josn}  # 传入到render变量参数用于被脚本调用
            return render(request, 'ScheduleParameter.html', context)


    