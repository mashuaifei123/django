import pandas as pd
from os import listdir
import os
import re
import zipfile
import copy
from siliao_1 import pyMuPDF2_fitz, pyMuPDF1_fitz
from  siliao import do1, do2
import sys, fitz
from os import listdir
import docx
from docx.shared import Inches

def del_with(roomid, start, end):
    start_y = start.split('-')[0]
    start_m = start.split('-')[1]
    start_d = start.split('-')[2]
    end_y = end.split('-')[0]
    end_m = end.split('-')[1]
    end_d = end.split('-')[2]

    start_month = str(int(start_m)) + '月'
    end_month = str(int(end_m)) + '月'
    print(start_month)
    print(end_month)
    if start_y == end_y:
        path_part2 = [start_y + '年度']
        month_list = [[str(i) + '月' for i in range(1, 13, 1)][int(start_m) - 1:int(end_m)]]
    else:
        path_part2 = [start_y+'年度',end_y+'年度']
        month_list = [[str(i) + '月' for i in range(1, 13, 1)][int(start_m)-1:], [str(i) + '月' for i in range(1, 13, 1)][: int(end_m)]]

    path_part1 = r'O:\共享资料\28.动物饲养相关信息\\'
    if len(path_part2) == 1:
        path = [path_part1 + path_part2[0]]
    else:
        path = [path_part1 + path_part2[i] for i in path_part2]

    for i in range(len(path)):
        docxlist = [fn for fn in listdir(path[i]) if fn.endswith('.xlsx') and fn.startswith('饲料') if fn[0] != '']
        bb = list(map(lambda y: path[i] + '\\' + y, docxlist))
        for j in bb:
            title = j[:-5].split('\\')[-1]
            print(title)
            df1 = pd.read_excel(j, header=None,sheet_name='1月至7月')
            df2 = pd.read_excel(j, header=None,sheet_name='8月-12月')
        #print(df1)
        siliao_total = []
        dianliao_total = []
        for m in range(len(df1[1])):
            c1 = ''.join(list(filter(str.isalnum, str(df1[1][m]))))
            c2 = ''.join(list(filter(lambda x:ord(x)<256,c1)))
            if c2 == roomid:
                index_c = m
        for m in range(len(df2[1])):
            c1 = ''.join(list(filter(str.isalnum, str(df2[1][m]))))
            c2 = ''.join(list(filter(lambda x:ord(x)<256,c1)))
            if c2 == roomid:
                index_c = m
        print(index_c)
        df_col_index = list(df1.iloc[0, :]) + list(df2.iloc[0, :])
        df_col = list(df1.iloc[index_c, :]) + list(df2.iloc[index_c, :])
        for n in month_list[i]:
            print(n)
            for nn in range(len(df_col_index)):
                if n in str(df_col_index[nn]):
                    df_dict = dict(zip(df_col_index[nn:nn+10], df_col[nn:nn+10]))
                    print(df_dict)
                    siliao = df_dict['饲料批号']
                    dianliao = df_dict['垫料批号']
                    beizhu = df_dict['备注']

                    if len(str(beizhu))> 5:
                        beizhu_list = list(filter(None, beizhu.split(' ')))

                        beizhu_value = [i.split('开')[0] if '开' in i else i.split('起')[0] for i in beizhu_list]
                        beizhu_key = [i[-8:] for i in beizhu_list]
                        beizhu_dict  = dict(zip(beizhu_key, beizhu_value))
                    else:
                        beizhu_dict ={}

                    if len(str(siliao)) > 5:
                        if len(str(siliao)) > 10:
                            siliao_list = list(filter(None, siliao.split(' ')))
                        else :
                            siliao_list = [siliao]
                    else:
                        siliao_list =[]

                    if len(str(dianliao)) > 5:
                        if len(str(dianliao)) > 10:
                            dianliao_list = list(filter(None, dianliao.split(' ')))
                        else :
                            dianliao_list = [dianliao]
                    else:
                        dianliao_list =[]
                    # print(siliao_list)
                    # print(dianliao_list)
                    # print(beizhu_dict)
        # 知道当月使用的饲料垫料之后，需要判断是否使用，存在
        # 当有两个以上饲料时，需要考虑某饲料是上个月用的延续，这时的时间点一般会在备注说明
        # 第一个月的判断
                    def catch_beizhu_day(dict_md):
                        if '月' in dict_md :
                            if '日' in dict_md:
                                day = (dict_md.split('日')[0]).split('月')[1]
                            else:
                                day = dict_md.split('月')[1]
                        else:
                            day = dict_md.split('.')[1]
                        return day

                    if n == start_month and len(siliao_list) > 2 :
                        dict_md = beizhu_dict[siliao_list[1]]
                        day = catch_beizhu_day(dict_md)
                        print(day)
                        if int(day) <= int(start_d):
                            siliao_list = siliao_list[2:]
                        else :
                            pass
                    if n == start_month and 3 > len(siliao_list) > 1:
                        dict_md = beizhu_dict[siliao_list[1]]
                        day = catch_beizhu_day(dict_md)
                        print(day)
                        if int(day) <= int(start_d):
                            siliao_list = siliao_list[2:]
                        else:
                            pass


                    if n == start_month and len(dianliao_list) > 1 :
                        dict_md = beizhu_dict[dianliao_list[1]]
                        day = catch_beizhu_day(dict_md)
                        if int(day) <= int(start_d):
                            dianliao_list = dianliao_list[1:]
                        else:
                            pass

            # 最后一个月的判断
                    if n == end_month and len(siliao_list) > 1:
                        if len(siliao_list) ==2:
                            dict_md = beizhu_dict[siliao_list[1]]
                            day = catch_beizhu_day(dict_md)
                            if int(day) <= int(end_d):
                                siliao_list = siliao_list[:2]
                            else:
                                siliao_list = siliao_list[0:1]
                        else:
                            dict_md1 = beizhu_dict[siliao_list[1]]
                            dict_md2 = beizhu_dict[siliao_list[2]]
                            day1 = catch_beizhu_day(dict_md1)
                            day2 = catch_beizhu_day(dict_md2)
                            if int(end_d)<= int(day2) :
                                if int(end_d)<= int(day1):
                                    siliao_list = siliao_list[0:1]
                                else :
                                    siliao_list = siliao_list[:2]
                            else:
                                pass

                    if n == end_month and len(dianliao_list) > 1:
                        if len(dianliao_list) ==2:
                            dict_md = beizhu_dict[dianliao_list[1]]
                            day = catch_beizhu_day(dict_md)
                            if int(day) <= int(end_d):
                                dianliao_list = dianliao_list[:2]
                            else:
                                dianliao_list = dianliao_list[0:1]
                        else:
                            dict_md1 = beizhu_dict[dianliao_list[1]]
                            dict_md2 = beizhu_dict[dianliao_list[2]]
                            day1 = catch_beizhu_day(dict_md1)
                            day2 = catch_beizhu_day(dict_md2)
                            if int(end_d)<= int(day2) :
                                if int(end_d)<= int(day1):
                                    dianliao_list = dianliao_list[0:1]
                                else :
                                    dianliao_list = dianliao_list[:2]
                            else:
                                pass

            if len(siliao_list)>0:
                siliao_total.append(siliao_list)
            if len(dianliao_list)>0:
                dianliao_total.append(dianliao_list)

        siliao_total1 = sum(siliao_total, [])
        dianliao_total1 = sum(dianliao_total, [])
        siliao_total = list(set(siliao_total1))
        dianliao_total = list(set(dianliao_total1))
        siliao_total.sort(key =siliao_total1.index)
        dianliao_total.sort(key =dianliao_total1.index)
        print(siliao_total)
        print(dianliao_total)
        # print(len(df_col_index))
        # print(len(df_col))
        return siliao_total, dianliao_total


def find_files(siliao,dianliao):
    # year = '20'+str(i)[:2]
    # path_bg = path1 + year + r'\\'
    # path_hgz = path2 + year + r'年度\饲料质量合格证\\'
    # pdflist_bg = [fn for fn in listdir(path_bg) if fn.startswith(str(i))][0]
    # animal_list = ['鼠', '猴', '犬']
    # #pdflist_hgz = [fn for fn in listdir(path_hgz) if fn.startswith(str(i))]
    # print(pdflist_bg)
    # #print(pdflist_hgz)
    # # 在路径里匹配文件后，pdf转
    # path_bg_file = path_bg + str(i)
    # path_hgz_file = path_hgz + str(i)

    """以上拼接路径作废"""
    path_bg = r'K:\mashuaifei\饲料查询\2020\\'
    path_hgz = r'K:\mashuaifei\饲料查询\饲料质量合格证\总\\'
    # 饲料
    bg_file = [fn for i in siliao for fn in listdir(path_bg) if fn.startswith(str(i))]  #文件夹
    path_bg_file = list(map(lambda y: path_bg + y, bg_file))  #文件夹路径
    path_bg_file_total = []
    for bg in path_bg_file:
        bg_pdf_file = [fn for fn in listdir(bg)] # 每个文件夹里的pdf名字
        path_bg_pdf_file = list(map(lambda y: bg + '\\' + y, bg_pdf_file)) # 每个文件夹下pdf的路径
        for file in path_bg_pdf_file:
            path_bg_file_total.append(file)
    hgz_file = [fn for i in siliao for fn in listdir(path_hgz) if fn.startswith(str(i))]
    path_hgz_file = list(map(lambda y: path_hgz + y, hgz_file))
    # 垫料
    bg_file1 = [fn for i in dianliao for fn in listdir(path_bg) if fn.startswith(str(i))]
    hgz_file1 = [fn for i in dianliao for fn in listdir(path_hgz) if fn.startswith(str(i))]
    print(path_bg_file_total)
    print(path_hgz_file)
    return path_bg_file_total,path_hgz_file


def finish(path1,path2,image_path1,image_path2, image_path3):
    image_path_list1 =[]
    image_path_list2 =[]
    print(path1)
    name = os.path.split(path1[:-4])[1]
    path1, path2 = find_files(siliao_total, dianliao_total)
    print(path1)
    print(path2)
    for i in path2:
        # print(i)
        pyMuPDF2_fitz(i, image_path2)
    for i in path1:
        # print(i)
        pyMuPDF2_fitz(i, image_path1)
        image_path_list1.append(image_path1 + '/' + name + '.png')
    for i in path1:
        # print(i)
        pyMuPDF1_fitz(i, image_path3)
        image_path_list2.append(image_path3 + '/' + name + '.png')
    doc2 = docx.Document()
    for i in image_path_list1:
        doc2.add_picture(i, width=docx.shared.Inches(5))
    doc2.save(image_path1 + '/' + name + '.docx')
    doc3 = docx.Document()
    for i in image_path_list2:
        doc3.add_picture(i, width=docx.shared.Inches(5))
    doc3.save(image_path3 + '/' + name + '.docx')


    dict1 = do1(image_path2)
    dict2 = do2(image_path1)
    dict3 = do2(image_path3)
    print(dict1)
    print(dict2)
    print(dict3)

if __name__ == "__main__":
    # time_start = time.clock()
    pdf_path1 = r'O:\共享资料\29.SD项目信息核对\4.兽医运行部\动物饲料检测报告\\'
    pdf_path2 = r'O:\共享资料\28.动物饲养相关信息\\ '
    image_path1 = r'K:\mashuaifei\饲料查询\图片1\\'
    image_path2 = r'K:\mashuaifei\饲料查询\图片2\\'
    image_path3 = r'K:\mashuaifei\饲料查询\图片3\\'
    roomid ='1351'
    time_strat = '2020-10-30'
    time_end = '2020-12-03'
    siliao_total, dianliao_total =  del_with(roomid, time_strat, time_end)
    finish(siliao_total, dianliao_total, image_path1, image_path2, image_path3)
    #path1, path2 = find_files(siliao_total,dianliao_total,pdf_path1, pdf_path2)
