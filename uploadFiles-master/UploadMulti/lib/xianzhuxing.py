# -*- coding: utf-8 -*-

import os
import numpy as np
import xlrd
import pandas as pd
from os import listdir  # 设置虚拟路径,以供将py文件所在路径的类型文件全部导入
import win32com
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 段落,表格对齐方式位置模块
from docx.enum.section import WD_SECTION  # 增加新章节 NEW_PAGE 下一页
from docx.oxml.ns import qn  # 设置中文字体类型
from docx.shared import Cm, Pt  # 设置字体大小Pt磅值,设置行高Cm厘米
from win32com.client import Dispatch, constants
import time
from docx import Document
from docx.enum.text import WD_LINE_SPACING
import  re



def Route_excel(path_list):
    xlsx_list = [ fn for fn in listdir(path_list) if fn.endswith('.xlsx')]
    rule = {'体温雌': 0, '体温雄': 1, '呼吸雌': 2, '呼吸雄': 3, '心电雌': 4, '心电雄': 5, '血压雌': 6, '血压雄': 7}
    # 如有需要，就可以添加新的顺序。
    newList = sorted(xlsx_list ,key = lambda x: rule[x[:3]])
    bb = list(map(lambda y: path_list + y, newList))
    print(newList)
    return  bb


def Conf_word(path_word:str, Row_n:list, time:list, sig_l:list, table_NO_szie_start=1, table_NO_szie_stop=10):
    '''

    :param path_word:  the path of word
    :param Row_n:  A list of the number of lines per Excel. such as: [5, 7, 10, 14]
    :param time: Time of each index in Excel . Time is often the number of lines of body temperature such as :[5,7,5,7]
    :param table_NO_szie_start: Starting position of table in word
    :param table_NO_szie_stop: Ending position of table in word
    :return:  a Format word file

    '''


    for tbt in range(table_NO_szie_start, table_NO_szie_stop):
        print(tbt)
        try:
            docTBTOP = doc.Tables(tbt)
            docTBTOP.Borders.Enable = False  # 除去所有线框
            docTBTOP.Rows(1).Borders(-1).LineStyle = 1  # 加首行粗线
            docTBTOP.Rows(1).Borders(-1).LineWidth = 12
            docTBTOP.Rows(3).Borders(-1).LineStyle = 1  # 加指标粗线
            docTBTOP.Rows(3).Borders(-1).LineWidth = 12
            docTBTOP.Rows(Row_n[tbt-1] + 2).Borders(-3).LineStyle = 1  # 加末行粗线
            docTBTOP.Rows(Row_n[tbt-1] + 2).Borders(-3).LineWidth = 12
            for i in range(1,Row_n[tbt-1]):
                docTBTOP.Rows(i+3).Borders(-1).LineStyle = 2  # 加表格虚线
                docTBTOP.Rows(i+3).Borders(-1).LineWidth = 6
            docTBTOP.range.ParagraphFormat.Alignment = 1      # 设定所有表格中文字段落对齐属性.
            docTBTOP.range.cells.VerticalAlignment = 1

            for s in sig_l[tbt-1]:  # 单元格加粗
                print(s)
                a = s[0]
                b = s[1]
                docTBTOP.Cell(a,b).Range.bold = True

            '''
            When the variable of indicator is more  than 2
            Tables need to be lined to separate different variables.
            '''
            if Row_n[tbt-1] > time[tbt-1]:
                for j in range(time[tbt-1], Row_n[tbt - 1] ):
                    if j % time[tbt - 1] == 0:
                        docTBTOP.Rows(j+3).Borders(-1).LineStyle = 1  # 加表格细线　
                        docTBTOP.Rows(j+3).Borders(-1).LineWidth = 8

                for h in range(1, Row_n[tbt-1] + 2):
                    if (h  % time[tbt-1] )== 0:
                        times=int(h/time[tbt-1])
                        docTBTOP.Cell(3+(times-1)*time[tbt-1], 1).Merge(docTBTOP.Cell(2 + times*time[tbt-1], 1))
                cell_1 = docTBTOP.Cell(1, 2)
                cell_2 = docTBTOP.Cell(2, 2)
                cell_1.Merge(cell_2)
            cell_3 = docTBTOP.Cell(1, 1)
            cell_4 = docTBTOP.Cell(2, 1)
            cell_3.Merge(cell_4)

        except IndexError as e:
            print('配置word格式出现警告请联系管理员', e)
            pass

# 小组名字,表格,数据列表,新增参数: 试验日期,参数
def write(group_name :list  ,table ,excel,additional: int ,sig_tol:list):
    '''

    :param group_name: such as : group_name=['辅料对照组','GR1501低剂量组','GR1501中剂量组','GR1501高剂量组'], the same to group_name_unit
    :param table:  just a table
    :param excel:  just a excel
    :param additional:  Extra lines
    :param sig_total: a list about how to bold the cell.such as :[[3, 11], [2, 11], [4, 8]]
    :return:  a table fill with data
    '''
    #print(sig_tol)
    for j in range(len(group_name)):
        table.cell(0, j + additional).text = group_name[j]
        table.cell(1, j + additional).text = group_name_unit[j]
    wb = xlrd.open_workbook(excel)
    sheet_c = wb.sheet_by_index(0)
    for h in range(0, sheet_c.nrows):
        table.rows[h + 2].height = Cm(0.6)
        row_date = sheet_c.row_values(h)
        for k in range(df.shape[1]):
            table.cell(h + 2, k).text = row_date[k]
    return  table

# 调整表格格式,
def adjust(table):
    table.rows[0].height = Cm(0.7)
    table.rows[1].height = Cm(0.7)
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10.5)


def add_Transverse_section():
    '''
    Add a horizontal page
    notes：If the previous page is horizontal, the next page will be vertical after calling

    :return:
    '''
    new_section = meansd.add_section(start_type=WD_SECTION.NEW_PAGE)
    new_width, new_height = new_section.page_height, new_section.page_width
    new_section.page_width = new_width
    new_section.page_height = new_height


def add_Vertical_section():
    '''
    Add a vertical page
    notes：If the previous page is horizontal, the next page will be the same after calling

    :return:
    '''
    new_section = meansd.add_section(start_type=WD_SECTION.NEW_PAGE)


def add_head(a:str ,b:str) :
    '''
    :param a: such as '表五'，'表六'
    :param b: such ad '体温统计数据'
    :return:  a word's head title ,such as '表五 食蟹猴皮下注射给予GR1501 26周重复给药毒性试验 体温统计数据'

    '''
    head1 = meansd.add_paragraph(a + study_name)
    head1_1 = meansd.add_paragraph(b)
    head1.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    head1.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
    head1_1.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_Atable(df:pd.DataFrame, Var:str,TableName:str ,sig_list:list) :
    '''

    :param df: the excel nead to be writered into word,and get its lwngth .
    :param Var : Number_variables ,need to get group_time
    :param TableName:  the new table's name you create .such as :'体温(℃)统计汇总（mean±SD）Females 雌性'.

    :param sig_total: a list about how to bold the cell.such as :[[3, 11], [2, 11], [4, 8]]
    :return: a table
    '''
    #print(sig_list)
    sig_list1=sig_transform(sig_list)
    table_length.append(df.shape[0])
    group_time.append(int(df.shape[0] / Number_variables[Var]))
    table_name = meansd.add_paragraph(TableName)
    table_name.style.font.size = Pt(12)
    table_name.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
    table_name.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    table = meansd.add_table(rows=df.shape[0] + 2, cols=df.shape[1], style='Table Grid')

    if Var=='体温':
        table.cell(0, 0).text = '试验日期'
        table.cell(1, 0).text = ''
        table = write(group_name, table, i, 1,sig_list1)
    else:
        table.cell(0, 0).text = Var +'参数'
        table.cell(1, 0).text = ''
        table.cell(0, 1).text = '试验日期'
        table.cell(1, 1).text = ''
        table = write(group_name, table, i, 2,sig_list1)
    adjust(table)
    note = meansd.add_paragraph()  # 增加一个段落
    note1 = note.add_run(list(sig_list)[-1])
    note1.font.size = Pt(10.5)  # 8=小5号字体



def add_Tagging(document):
    tag1= '注：N：动物数；N＜3时不进行统计'
    tag = document.add_paragraph()


def sig_transform(ser:list ) -> list:
    '''
    sig1 : index length less than 10
    sig2 :  index length more than 10
    sig1= [[6, 2, 2]]
    sig2= [[11, 2, 2]]
    sig_total=[[7,11,2],[12,11,2]]
    Get a position to mark
    '''
    sig1=[]
    sig2=[]
    sig_total=[]
    length=0
    for ss in ser:
        if ss != '':
            length+=1
    for i in range(1,length-1):
        a = re.sub("[^0-9]", "", ser[i])
        if len(a)>3:
            b2=[ int(i) for i in [a[0]+a[1],a[2]]]
            sig2.append(b2)
        else:
            c2=[ int(i) for i in [a[0],a[1]]]
            sig1.append(c2)
    if sig1:
        sig1=[[j[0]+2,j[1]+2]  for j in sig1]
        for s in sig1:
            sig_total.append(s)
    if sig2:
        sig2=[[j[0]+2,j[1]+2]  for j in sig2]
        for s in sig2:
            sig_total.append(s)
    #print(sig_total)
    return sig_total




#if __name__ == "__main__":


start = time.time()
special_topic='A2019001-T014-01'
study_name = '食蟹猴皮下注射给予GR1501 26周重复给药毒性试验'
group_name=['辅料对照组','GR1501低剂量组','GR1501中剂量组','GR1501高剂量组']
group_name_unit=['（0 mg/kg）','（15 mg/kg）','（50 mg/kg）','（150 mg/kg）']
Number_variables = {'体温': 1, '呼吸': 2, '心电': 9, '血压': 3} # 变量的数量，如有需要，需要新添加

group_time=[]
path_list=r'K:\mashuaifei\新建文件夹\新建文件夹\高婷婷他4\新建文件夹 (2)\\'
word_path_moban = r'K:\mashuaifei\显著性判断\word\个人实验.docx'

bb= Route_excel(path_list)
meansd = Document(word_path_moban)
meansd.styles['Normal'].font.name = u'宋体'
meansd.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
meansd.styles['Normal'].font.name = u'Times New Roman'
meansd.styles['Normal']._element.rPr.rFonts.set(qn('w:westAsia'), u'Times New Roman')
NO_end = 1
table_length = []
flag = 1
ConfigDf=pd.read_excel('K:\mashuaifei\新建文件夹\新建文件夹\高婷婷他4\config.xlsx',header=None)
ConfigDf.replace(np.nan,'',inplace=True)
Sigtotal=[]


for i in bb:
    df = pd.read_excel(i,header=None)
    if '体温' in i  :
        if flag == 1:
            add_head('表五','体温统计数据')
            flag = 0
        if '雌' in i:
            add_Transverse_section()
            add_Atable(df,'体温','体温(℃)统计汇总（mean±SD）Females 雌性',ConfigDf[0])
            sig_list = sig_transform(ConfigDf[0])
            Sigtotal.append(sig_list)
            NO_end+=1

        if '雄' in i:
            add_Vertical_section()
            add_Atable(df, '体温', '体温(℃)统计汇总（mean±SD）Females 雄性',ConfigDf[1])
            sig_list = sig_transform(ConfigDf[1])
            Sigtotal.append(sig_list)
            NO_end+=1
            flag=1

    if '呼吸' in i:
        if flag == 1:
            add_Transverse_section()
            add_head('表六','呼吸统计数据')
            flag = 0
        if '雌' in i:
            add_Transverse_section()
            add_Atable(df, '呼吸', '呼吸统计汇总（mean±SD）Females 雌性',ConfigDf[2])
            sig_list = sig_transform(ConfigDf[2])
            Sigtotal.append(sig_list)
            NO_end+=1
        if '雄' in i:
            add_Vertical_section()
            add_Atable(df, '呼吸', '呼吸统计汇总（mean±SD）Females 雌性',ConfigDf[3])
            sig_list = sig_transform(ConfigDf[3])
            Sigtotal.append(sig_list)
            NO_end+=1
            flag = 1

    if '心电' in i:
        if flag == 1:
            add_Transverse_section()
            add_head('表七','心电统计数据')
            flag = 0
        if '雌' in i:
            add_Transverse_section()
            add_Atable(df, '心电', '心电统计汇总（mean±SD）Females 雌性',ConfigDf[4])
            sig_list = sig_transform(ConfigDf[4])
            Sigtotal.append(sig_list)
            NO_end+=1
        if '雄' in i:
            add_Vertical_section()
            add_Atable(df, '心电', '心电统计汇总（mean±SD）Females 雌性',ConfigDf[5])
            sig_list = sig_transform(ConfigDf[5])
            Sigtotal.append(sig_list)
            NO_end+=1
            flag = 1

    if '血压' in i:
        if flag == 1:
            add_Transverse_section()
            add_head('表八','血压统计数据')
            flag = 0
        if '雌' in i:
            add_Transverse_section()
            add_Atable(df, '血压', '血压统计汇总（mean±SD）Females 雌性',ConfigDf[6])
            sig_list = sig_transform(ConfigDf[6])
            Sigtotal.append(sig_list)
            NO_end+=1
        if '雄' in i:
            add_Vertical_section()
            add_Atable(df, '血压', '血压统计汇总（mean±SD）Females 雌性',ConfigDf[7])
            sig_list = sig_transform(ConfigDf[7])
            Sigtotal.append(sig_list)
            NO_end+=1
            flag = 1


NO_start = 1
meansd.save(path_list+r'\demo.docx')
time.sleep(1)
w = win32com.client.Dispatch('Word.Application')
path_word = path_list + r'\demo.docx'
doc = w.Documents.Open(path_word)
#myRange = doc.Range(doc.Content.Start, doc.Content.End)  # 全选
#myRange.Style.Font.Size = 10.5
#pf = myRange.ParagraphFormat
#pf.LineSpacingRule = 1
#print(table_length)
#print(group_time)
Conf_word(path_word, table_length, group_time, Sigtotal, NO_start, NO_end)
doc.Save()
doc.Close()
end = time.time()
