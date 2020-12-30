import os # 设置虚拟路径,以供将py文件所在路径的类型文件全部导入
from os import listdir
import numpy as np
import pandas as pd
import win32com.client
from docx import Document  # docx用来操作word.docx,document用来新建空白文档
from docx.enum.text import WD_LINE_SPACING  # 设置段落的行间距
from docx.oxml.ns import qn  # 设置中文字体类型
from docx.shared import Cm, Pt  # 设置字体大小Pt磅值,设置行高Cm厘米
import time
import re
import pythoncom
import docx
import re


class RemoteWord:
    def __init__(self, filename=None):
        self.xlApp = win32com.client.DispatchEx('Word.Application')
        self.xlApp.Visible = 0
        self.xlApp.DisplayAlerts = 0  # 后台运行，不显示，不警告
        if filename:
            self.filename = filename
            if os.path.exists(self.filename):
                file_pass = True
                jishu = 0
                while file_pass:
                    try:
                        self.doc = self.xlApp.Documents.Open(filename, False, True)  # 第2个True只读打开文件.
                        file_pass = False
                    except:
                        jishu += 1
                        time.sleep(10)
                        if jishu == 5:
                            print('程序出现错误.请定位到文件打开部分.35行')
                            file_pass = False
            else:
                self.doc = self.xlApp.Documents.Add()  # 创建新的文档
                self.doc.SaveAs(filename)
        else:
            self.doc = self.xlApp.Documents.Add()
            self.filename = ''

    def add_doc_end(self, string):
        '''在文档末尾添加内容'''
        rangee = self.doc.Range()
        rangee.InsertAfter('\n' + string)

    def add_doc_start(self, string):
        '''在文档开头添加内容'''
        rangee = self.doc.Range(0, 0)
        rangee.InsertBefore(string + '\n')

    def insert_doc(self, insertPos, string):
        '''在文档insertPos位置添加内容'''
        rangee = self.doc.Range(0, insertPos)
        if (insertPos == 0):
            rangee.InsertAfter(string)
        else:
            rangee.InsertAfter('\n' + string)

    def replace_doc(self, string, new_string):
        '''替换文字'''
        self.xlApp.Selection.Find.ClearFormatting()
        self.xlApp.Selection.Find.Replacement.ClearFormatting()
        self.xlApp.Selection.Find.Execute(string, False, False, False, False, False, True, 1, True, new_string, 2)

    def save(self):
        '''保存文档'''
        self.doc.Save()

    def save_as(self, filename):
        '''文档另存为'''
        self.doc.SaveAs(filename)

    def w_to_pdf(self, out_path: str):
        '''另存为pdf'''
        self.doc.SaveAs2(out_path, 17)
        self.xlApp.Documents.Close()

    def close(self):
        '''保存文件、关闭文件'''
        # self.save()
        self.xlApp.Documents.Close()
        self.xlApp.Quit()

    def Quit(self):
        '''退出出word'''
        self.xlApp.Quit()

    def PageSetup_Orientation(self, Find_string) -> int:
        '''查找文本内容所在章节或者页面并返回章节'''
        self.xlApp.Selection.Find.ClearFormatting()
        #  查找字符串查找完代表选中 注意匹配模式全文搜内容第78个布尔值
        self.xlApp.Selection.Find.Execute(Find_string, False, False, False, False, False, True, 1, True)
        # 可以返回查找文件所在章节
        page_number = self.xlApp.Selection.Information(2)
        return page_number
        # 把选中的内容所在章节设定为纵向
        # doc_s5 = self.xlApp.Selection.Sections[page_number]
        # PageSetup.PaperSize
        # doc_s5.PageSetup.Orientation = 0
        # self.xlApp.Selection.Sections[4].PageSetup.Orientation = 0


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None


def word_to_p(word_path, out_path):
    pythoncom.CoInitialize()
    docxlist = [fn for fn in listdir(word_path) if fn.endswith('.docx') if fn[0] != '~']
    bb = list(map(lambda y: word_path + '\\' + y, docxlist))
    for j in bb:
        print(bb)

        dpos_n_l = set()
        listfind = ['Individual Animal Report of  Daily Clinical Signs Days Seen'
        ]
        doc = RemoteWord(j)  # 初始化一个doc对象 win32aip
        for o in listfind:
            Selec_info = doc.PageSetup_Orientation(o)
            dpos_n_l.add(Selec_info)
        doc.close()
        print(dpos_n_l)
        file = docx.Document(j)
        dsec_len = len(file.sections)  # word共有多少章节
        print(dsec_len)
        for i in dpos_n_l:  # 循环一个找到的节 集合设定页面方向
            try:
                section = file.sections[i-1]
                section.header_distance = Cm(0.85)
                new_width, new_height = section.page_height, section.page_width
                section.page_width, section.page_height  = new_width, new_height
            except:
                print('转换方向错误')

        # for paragraph in file.paragraphs:
        #     print(paragraph.text)
        lines = [0 for i in range(len(file.paragraphs))]

        k = 0
        for paragraph in file.paragraphs:
            lines[k] = paragraph.text
            k = k + 1
        # print(lines)
        appendix_list = ['P = Pretest Phase', 'PreDosing(F) = PreDosing(F)', 'PreDosing(M) = PreDosing(M)',
                         'D = Dosing Phase', 'R = Recovery Phase', 'PreD= PreDosing', '“”表示不适用', 'L=耗食减少',
                         'N=耗食正常', 'Z=未摄食', '“”代表阴性。','P=Pretest Phase            R=Recovery Phase     ',
                         'P=Pretest Phase            PreD=PreDosing Phase            D=Dosing Phase            R=Recovery Phase     ',
                         'P=Pretest Phase            D=Dosing Phase            R=Recovery Phase     ',
                         'L=耗食减少      N=耗食正常      Z=未摄食     ',
                         'Y=YELLOW    LO=LIGHTORANGE    O=ORANGE    LY=LIGHTYELLOW    CL=COLORLESS    LB=LIGHTBROWN   ',
                         'P=Pretest Phase            D=Dosing Phase            R=Recovery Phase     ',
                         'R=Recovery Phase     '
                         ]

        find_list_total = ['一般状态观察个体数据',
                           '体重测定结果个体数据（kg）',
                           '摄食量检测结果个体数据（kg）',
                           '摄食量测定结果个体数据',
                           '耗食量个体数据',
                           '血液学指标检测结果个体数据',
                           '凝血指标检测结果个体数据',
                           '血生化指标检测结果个体数据',
                           '尿液指标测定结果个体数据',
                           '尿沉渣指标检查结果个体数据',
                           '脏器重量（g）个体数据',
                           '脏体比（%）计算结果个体数据',
                           '脏脑比（%）计算结果个体数据'

                          ]
        # 筛选此word含有的表格名字
        find_list1 = [j for j in find_list_total for i in lines if ''.join(re.findall('[\u4e00-\u9fa5]+', i)) == ''.join(re.findall('[\u4e00-\u9fa5]+', j))]
        find_list = list(set(find_list1))
        find_list.sort(key =find_list1.index)
        print(find_list)
        # 找到 两个关键分隔点的位置
        remove_list = ['Report Selections']
        remove_location =[i for i in range(len(lines)) if lines[i].strip() == remove_list[0] ]
        cease_list = ['\n']
        cease_location = [i for i in range(len(lines)) if lines[i] == cease_list[0] ]
        # print(remove_location)
        # print(cease_location)

        # for num, paragraphs in enumerate(file.paragraphs):
        #     print(paragraphs)
        #     if len(paragraphs.text.strip()) == 0:
        #         print(第{}段是空行段.format(num))
        #         paragraphs.clear()  # 清除文字，并不删除段落，run也可以,paragraph.run.clear()
        #         del paragraphs
        #         # delete_paragraph(paragraphs)


        # 找到具体的表格名字位置
        for jj in range (len(find_list)):
            locations = []
            for i in range(len(lines)):
                if ''.join(re.findall('[\u4e00-\u9fa5]+', lines[i])) ==''.join(re.findall('[\u4e00-\u9fa5]+', find_list[jj])) :
                    locations.append(i)
            locations.pop(0)
            start_location = locations
            end_location = [l-2 for l in locations[1:]] + [remove_location[jj]]
            print(start_location,end_location)
            if len(start_location)== len(end_location):
                # for p in range(len(start_location)):
                #     page = lines[start_location[p]:end_location[p]]
                #     page_list = [n for n in appendix_list if n in page]
                #     print(page_list)
                page = lines[start_location[0]:end_location[0]]
                page_list = [n for n in appendix_list if n in page]
                print(page_list)
                # print(lines.index('P = Pretest Phase'))
                if '一般状态' in find_list[jj]:
                    print('gggg')
                    str1 = 'Note：P = Pretest Phase     PreD= Pre-Dosing     D = Dosing Phase     R = Recovery Phase'
                    file.paragraphs[lines.index('P = Pretest Phase')].text = ''
                    file.paragraphs[lines.index('P = Pretest Phase')].insert_paragraph_before(str1)
                    file.paragraphs[lines.index('PreDosing(F) = PreDosing(F)')].clear()
                    del file.paragraphs[lines.index('PreDosing(F) = PreDosing(F)')]
                    file.paragraphs[lines.index('PreDosing(M) = PreDosing(M)')].clear()
                    del file.paragraphs[lines.index('PreDosing(M) = PreDosing(M)')]
                    file.paragraphs[lines.index('D = Dosing Phase')].text =''
                    file.paragraphs[lines.index('R = Recovery Phase')].text = ''
        #         if len(page_list)==1:
        #             if page_list[0] =='R=Recovery Phase     ':
        #                 strs = 'Note：R=Recovery Phase'
        #                 strs1 = 'HT=Heart    LWG=Liver/gall bladder    SP=Spleen    LBR=Lung/main bronchi    KD=Kidney    BRN=Brain    ' \
        #                         ' TMA=Thymus or thymic area GAD=Gland, adrenal    GPI=Gland, pituitary    GTP=Gland, thyroid/gland, parathyroid    OV=Ovary    ' \
        #                         'UC=Uterus/cervix uteri    TE=Testis    EPI=Epididymis    GPS=Gland, prostate'
        #                 file.paragraphs[lines.index('R=Recovery Phase     ')].text =strs1
        #                 file.paragraphs[lines.index('R=Recovery Phase     ')].insert_paragraph_before(strs)
        #             elif page_list[0] =='P=Pretest Phase            R=Recovery Phase     ':
        #                 strs = 'Note：P=Pretest Phase            R=Recovery Phase     '
        #                 file.paragraphs[lines.index('R=Recovery Phase     ')].text = strs
        #         elif len(page_list) ==2:
        #             pass
        #
        print(lines)
        # # print(len(lines))
        # print(file.paragraphs[3].text)




        file.save(r'K:\mashuaifei\pristima_wp\temp.docx')


if __name__ == "__main__":
    # time_start = time.clock()
    word_path = r'K:\mashuaifei\pristima_wp\1'
    out_path = r'K:\mashuaifei\pristima_wp\1'
    word_to_p(word_path, out_path)
    # print('恭喜你搞定了总耗时', time.clock() - time_start)